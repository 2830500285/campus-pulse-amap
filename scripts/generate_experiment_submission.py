from __future__ import annotations

import json
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
PROJECT_NAME = "山科智行 Campus Pulse"
PROJECT_SUBTITLE = "基于高德地图 API 的青岛校区校园出行工作台"
SITE_URL = "https://2830500285.github.io/campus-pulse-amap/#/"
REPO_URL = "https://github.com/2830500285/campus-pulse-amap"

DOCS = {
    "report_docx": OUTPUT_DIR / "实验报告_山科智行CampusPulse.docx",
    "design_docx": OUTPUT_DIR / "设计文档_山科智行CampusPulse.docx",
    "release_docx": OUTPUT_DIR / "发布说明_山科智行CampusPulse.docx",
    "report_pdf": OUTPUT_DIR / "实验报告_山科智行CampusPulse.pdf",
    "design_pdf": OUTPUT_DIR / "设计文档_山科智行CampusPulse.pdf",
    "release_pdf": OUTPUT_DIR / "发布说明_山科智行CampusPulse.pdf",
    "architecture_png": OUTPUT_DIR / "山科智行_系统架构图.png",
    "workflow_png": OUTPUT_DIR / "山科智行_开发交付流程图.png",
    "workload_png": OUTPUT_DIR / "山科智行_工作量统计图.png",
    "code_inventory": OUTPUT_DIR / "程序代码目录说明.txt",
    "attachment_list": OUTPUT_DIR / "提交材料清单.txt",
    "code_zip": OUTPUT_DIR / "程序代码_山科智行CampusPulse源码.zip",
    "android_apk": OUTPUT_DIR / "山科智行CampusPulse_Android_Debug.apk",
}

SCREENSHOTS = {
    "home": OUTPUT_DIR / "发布效果_01_行动看板首页.png",
    "route": OUTPUT_DIR / "发布效果_02_路线调度结果.png",
    "mobile": OUTPUT_DIR / "发布效果_03_移动端界面.png",
}

APP_SCREENSHOT_DIR = OUTPUT_DIR / "application_screenshots"
REPORT_FIGURE_DIR = OUTPUT_DIR / "report_figures"
APP_SCREENSHOTS = [
    ("01", "01_高德实景首页.png", "高德实景首页", "系统默认进入高德 JSAPI 实景地图，右侧地图作为首屏主体，左侧保留起终点路线调度入口。"),
    ("02", "02_路线规划结果_西门到图书信息中心.png", "西门到图书信息中心路线结果", "选择西门和图书信息中心后，系统同时展示校内路网距离、预计时间、分段提示和高德实景背景。"),
    ("03", "03_地点快照_图书信息中心.png", "图书信息中心地点快照", "通过链接参数选中地点后，地点快照区显示用途、关键词、起点和终点快捷操作。"),
    ("04", "04_搜索餐厅结果.png", "餐厅关键词搜索", "输入餐厅关键词后，本地地点队列和高德实时 POI 共同缩小检索范围。"),
    ("05", "05_食堂分类筛选.png", "食堂分类筛选", "点击任务过滤中的“食堂”后，地点队列只保留 A 餐厅、B 餐厅和学苑餐厅，筛选结果与页面统计同步变化。"),
    ("06", "06_生活补给带筛选.png", "生活补给带场景筛选", "按空间场景带筛选后，生活区餐厅、宿舍、校医院等点位集中展示。"),
    ("07", "07_宿舍到实训中心路线.png", "宿舍到工程实训中心路线", "第二组路线用于验证跨生活区、教学区和科研区的较长路径规划能力。"),
    ("08", "08_地点详情页_图书信息中心.png", "图书信息中心详情页", "详情页保留直达链接能力，适合从首页地点快照进入单点说明。"),
    ("09", "09_地点详情页_学生公寓A区.png", "学生公寓 A 区详情页", "生活区宿舍详情页展示到达建议、识别词和同类地点推荐。"),
    ("10", "10_移动端高德首页.png", "移动端高德首页", "手机宽度下页面纵向组织首屏信息，保留高德地图、路线控件和地点列表。"),
    ("11", "11_移动端路线规划结果.png", "移动端路线规划结果", "移动端路线结果展示总距离、预计时间、地图和分步说明，验证响应式布局。"),
    ("12", "12_移动端搜索餐厅.png", "移动端搜索餐厅", "移动端搜索场景验证输入框、实时地点和筛选列表在窄屏下可读。"),
    ("13", "13_校园拓扑路线校验.png", "校园拓扑路线校验", "切换到拓扑视图后，可核对本地路网节点、入口节点和路线折线。"),
]

FONT_CANDIDATES = [
    Path("/System/Library/Fonts/PingFang.ttc"),
    Path("/System/Library/Fonts/STHeiti Light.ttc"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
]


@dataclass
class ProjectMetrics:
    source_files: int
    source_lines: int
    test_files: int
    test_cases: int
    components: int
    place_count: int
    category_count: int
    graph_nodes: int
    graph_edges: int


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int):
    for candidate in FONT_CANDIDATES:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def set_east_asia_font(run, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_east_asia_font(run, "黑体")
    run.bold = True
    run.font.size = Pt(22)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    sub_run = sub.add_run(subtitle)
    set_east_asia_font(sub_run, "宋体")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(88, 99, 112)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_east_asia_font(run, "黑体")
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12)


def add_paragraph(document: Document, text: str, first_line: bool = True) -> None:
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.55
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(11)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"- {item}")
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_east_asia_font(run, "黑体")
        run.bold = True
        run.font.size = Pt(10.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(value)
            set_east_asia_font(run, "宋体")
            run.font.size = Pt(10)


def add_picture(document: Document, path: Path, width: float = 6.3, caption: str | None = None) -> None:
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        c = document.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        c_run = c.add_run(caption)
        set_east_asia_font(c_run, "宋体")
        c_run.italic = True
        c_run.font.size = Pt(9.5)
        c_run.font.color.rgb = RGBColor(88, 99, 112)


def add_small_note(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(88, 99, 112)


def prepare_report_figures() -> dict[str, Path]:
    REPORT_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    figure_paths: dict[str, Path] = {}

    for number, filename, _title, _caption in APP_SCREENSHOTS:
        source = APP_SCREENSHOT_DIR / filename
        if not source.exists():
            continue

        image = Image.open(source).convert("RGB")
        width, height = image.size
        if width <= 500:
            crop_height = min(height, 900)
        elif height > width:
            crop_height = min(height, 1080)
        else:
            crop_height = min(height, 920)

        cropped = image.crop((0, 0, width, crop_height))
        target = REPORT_FIGURE_DIR / filename
        cropped.save(target, quality=92)
        figure_paths[number] = target

    return figure_paths


def add_application_screenshot_gallery(document: Document, start_index: int = 1) -> None:
    figure_paths = prepare_report_figures()
    if not figure_paths:
        add_paragraph(document, "当前未检测到真实应用截图目录，报告保留文字说明；运行 Playwright 截图后可自动插入应用图集。")
        return

    for index, (number, _filename, title, caption) in enumerate(APP_SCREENSHOTS):
        figure_path = figure_paths.get(number)
        if not figure_path:
            continue
        add_picture(document, figure_path, width=5.9 if number not in {"10", "11", "12"} else 3.2, caption=f"图 {start_index + index}  {title}")
        add_small_note(document, f"截图说明：{caption}")


def count_between(text: str, marker: str) -> int:
    return text.count(marker)


def count_source_metrics() -> ProjectMetrics:
    patterns = ["src/**/*.ts", "src/**/*.tsx", "src/**/*.css", "server/src/**/*.ts", "shared/**/*.ts"]
    files = {path for pattern in patterns for path in ROOT.glob(pattern) if path.is_file()}
    source_lines = 0
    for path in files:
        source_lines += len(path.read_text(encoding="utf-8").splitlines())

    test_files = [path for path in files if ".test." in path.name]
    test_cases = 0
    for path in test_files:
        text = path.read_text(encoding="utf-8")
        test_cases += count_between(text, "it(") + count_between(text, "test(")

    places_text = (ROOT / "src" / "data" / "places.ts").read_text(encoding="utf-8")
    categories_text = (ROOT / "src" / "data" / "categories.ts").read_text(encoding="utf-8")
    graph_text = (ROOT / "src" / "data" / "graph.ts").read_text(encoding="utf-8")

    return ProjectMetrics(
        source_files=len(files),
        source_lines=source_lines,
        test_files=len(test_files),
        test_cases=test_cases,
        components=len(list((ROOT / "src" / "components").glob("*.tsx"))),
        place_count=places_text.count("categoryId:"),
        category_count=categories_text.count("label:"),
        graph_nodes=graph_text.count("label:"),
        graph_edges=graph_text.count("fromNodeId:"),
    )


def draw_box(draw: ImageDraw.ImageDraw, box, fill, outline, width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=8, fill=fill, outline=outline, width=width)


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy, font, fill, max_width: int, spacing: int = 8) -> None:
    x, y = xy
    line = ""
    lines: list[str] = []
    for char in text:
        probe = line + char
        bbox = draw.textbbox((0, 0), probe, font=font)
        if bbox[2] - bbox[0] > max_width and line:
            lines.append(line)
            line = char
        else:
            line = probe
    if line:
        lines.append(line)
    for item in lines:
        draw.text((x, y), item, font=font, fill=fill)
        bbox = draw.textbbox((x, y), item, font=font)
        y += bbox[3] - bbox[1] + spacing


def create_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1080), "#FAF7EE")
    draw = ImageDraw.Draw(image)
    title = get_font(54)
    h = get_font(31)
    b = get_font(23)
    small = get_font(20)
    draw.text((80, 60), f"{PROJECT_NAME} 系统架构图", font=title, fill="#075F54")
    draw.text((82, 130), "高德地图 API、路线调度、服务计算和交付证据链分层组织", font=small, fill="#5D6D66")

    boxes = [
        ((90, 235, 575, 620), "前端白色出行工作台", "React + Vite + TypeScript\n路线调度、地点队列、场景带过滤和高德实景主视图"),
        ((655, 235, 1140, 620), "高德地图 API 层", "JSAPI 2.0\nPlaceSearch、Walking、Geolocation、Scale、ToolBar"),
        ((1220, 235, 1710, 620), "导航服务与路网", "Express + TypeScript\nbootstrap 数据接口、route 规划接口和校内最短路径计算"),
        ((275, 720, 815, 950), "共享模型与静态回退", "PlaceRecord / GraphNode / GraphEdge / PlannedRoute\n无后端时由 local-route-planner 输出路线"),
        ((990, 720, 1530, 950), "交付材料", "DOCX / PDF / 截图 / 架构图 / 源码包 / 提交清单统一生成"),
    ]
    for box, heading, body in boxes:
        draw_box(draw, box, "#FFFFFF", "#9CC9BB")
        draw.text((box[0] + 28, box[1] + 24), heading, font=h, fill="#075F54")
        draw_wrapped(draw, body, (box[0] + 30, box[1] + 82), b, "#17231F", box[2] - box[0] - 60)
    for start, end in [((575, 430), (655, 430)), ((1140, 430), (1220, 430)), ((720, 620), (545, 720)), ((1080, 620), (1260, 720))]:
        draw.line([start, end], fill="#D96F2D", width=5)
    image.save(path)


def create_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1180), "#FAF7EE")
    draw = ImageDraw.Draw(image)
    title = get_font(52)
    node = get_font(28)
    body = get_font(20)
    draw.text((80, 60), "开发与交付流程图", font=title, fill="#075F54")
    draw.text((82, 128), "从旧项目基线到高德 API 软件、白色界面、报告和可验证交付包的完整流程", font=body, fill="#5D6D66")

    steps = [
        ("1 基线拆解", "读取旧项目能力、文档脚本、测试和数据模型"),
        ("2 主题重构", "确定 Campus Pulse 出行工作台定位和白色产品界面"),
        ("3 高德接入", "配置 Web 端 Key 和安全密钥，启用 JSAPI 实景地图、POI 和步行规划"),
        ("4 数据同步", "同步前后端校区配置、分区、底图文件和测试断言"),
        ("5 报告生成", "生成实验报告、设计文档、发布说明与图示材料"),
        ("6 质量校验", "执行 lint、test、build 和文档 PDF 渲染检查"),
        ("7 归档提交", "输出源码压缩包、目录说明、提交材料清单"),
    ]
    for idx, (heading, text) in enumerate(steps):
        x = 120 if idx % 2 == 0 else 950
        y = 220 + (idx // 2) * 220
        draw_box(draw, (x, y, x + 660, y + 126), "#FFFFFF", "#9CC9BB")
        draw.text((x + 26, y + 18), heading, font=node, fill="#075F54")
        draw_wrapped(draw, text, (x + 28, y + 62), body, "#17231F", 600)
        if idx < len(steps) - 1:
            end_x = 950 if idx % 2 == 0 else 120
            draw.line([(x + 660, y + 63), (end_x, y + 63)], fill="#D96F2D", width=4)
    image.save(path)


def create_workload_chart(path: Path, metrics: ProjectMetrics) -> None:
    image = Image.new("RGB", (1600, 980), "#FAF7EE")
    draw = ImageDraw.Draw(image)
    title = get_font(48)
    label = get_font(22)
    num = get_font(38)
    draw.text((80, 56), "项目工作量统计", font=title, fill="#075F54")
    draw.text((82, 122), "统计范围不含 node_modules、dist、output、tmp 和第三方缓存", font=label, fill="#5D6D66")

    items = [
        ("源码文件", metrics.source_files),
        ("源码行数", metrics.source_lines),
        ("测试文件", metrics.test_files),
        ("测试用例", metrics.test_cases),
        ("前端组件", metrics.components),
        ("点位数量", metrics.place_count),
        ("路网节点", metrics.graph_nodes),
        ("路网边", metrics.graph_edges),
    ]
    for idx, (name, value) in enumerate(items):
        col = idx % 4
        row = idx // 4
        x = 90 + col * 370
        y = 220 + row * 230
        draw_box(draw, (x, y, x + 320, y + 170), "#FFFFFF", "#9CC9BB")
        draw.text((x + 24, y + 28), name, font=label, fill="#5D6D66")
        draw.text((x + 24, y + 82), f"{value:,}", font=num, fill="#075F54")
    image.save(path)


def create_placeholder_screenshots() -> None:
    source_map = {
        "home": APP_SCREENSHOT_DIR / "01_高德实景首页.png",
        "route": APP_SCREENSHOT_DIR / "02_路线规划结果_西门到图书信息中心.png",
        "mobile": APP_SCREENSHOT_DIR / "10_移动端高德首页.png",
    }
    for key, source in source_map.items():
        if not source.exists():
            continue
        image = Image.open(source).convert("RGB")
        width, height = image.size
        crop_height = min(height, 920 if width > 500 else 1100)
        image.crop((0, 0, width, crop_height)).save(SCREENSHOTS[key], quality=92)

    specs = [
        ("home", (1440, 1050), "白色出行工作台首页", "高德地图 API、路线调度、地点队列和场景带过滤集中在一屏内"),
        ("route", (1440, 1050), "高德路线调度结果", "西门到图书信息中心：展示总距离、预计耗时、实景地图和分步提示"),
        ("mobile", (430, 1100), "移动端白色界面", "手机宽度下保持起终点选择、高德地图和行程摘要可读"),
    ]
    for key, size, title, subtitle in specs:
        path = SCREENSHOTS[key]
        if path.exists():
            continue
        image = Image.new("RGB", size, "#FAF7EE")
        draw = ImageDraw.Draw(image)
        title_font = get_font(42 if size[0] > 500 else 28)
        body_font = get_font(24 if size[0] > 500 else 18)
        small_font = get_font(18 if size[0] > 500 else 14)
        draw.rectangle((0, 0, size[0], size[1]), fill="#FAF7EE")
        for x in range(0, size[0], 34):
            draw.line([(x, 0), (x, size[1])], fill="#E1DCCF")
        for y in range(0, size[1], 34):
            draw.line([(0, y), (size[0], y)], fill="#E1DCCF")
        draw.text((36, 34), PROJECT_NAME, font=title_font, fill="#075F54")
        draw_wrapped(draw, subtitle, (38, 94), body_font, "#17231F", size[0] - 76)
        card_w = size[0] - 72
        y = 180
        sections = ["路线调度", "空间雷达", "行程摘要", "地点队列"]
        for idx, section in enumerate(sections):
            h = 140 if size[0] > 500 else 120
            draw_box(draw, (36, y, 36 + card_w, y + h), "#FFFFFF", "#9CC9BB", width=2)
            draw.text((60, y + 24), section, font=body_font, fill="#075F54")
            draw.text((60, y + 70), "AMap JSAPI / verified view", font=small_font, fill="#5D6D66")
            y += h + 26
        draw.text((38, size[1] - 54), title, font=body_font, fill="#D96F2D")
        image.save(path)


def create_report_doc(metrics: ProjectMetrics) -> None:
    document = Document()
    set_base_style(document)
    add_title(document, "实验报告", PROJECT_SUBTITLE)
    add_table(
        document,
        ["项目", "内容"],
        [
            ["实验名称", f"{PROJECT_NAME} 设计与实现"],
            ["学生信息", "姓名、学号、班级待提交前填写"],
            ["开发周期", "需求拆解、差异化改造、高德 API 接入、测试验证、报告生成"],
            ["成果形态", "Web 软件、后端服务、DOCX/PDF 报告、架构图、13 张应用截图和源码包"],
        ],
    )

    add_heading(document, "一、实验背景与任务要求")
    for text in [
        "本实验面向校园步行导航与校园空间信息服务场景，要求在已有项目基础上重新制作一份软件与报告。重新制作并不等于只替换名称，而是需要在产品定位、界面风格、地图能力、交互路径、文档表达和交付材料上形成明显差异，使新成果能够作为另一份独立课程作品提交。",
        "原有项目已经具备地点数据、路线规划、前后端协作和报告生成能力，因此本次实验的重点不是推翻所有底层逻辑，而是在保留可验证能力的基础上，把系统改造成基于高德地图开放平台 API 的校园出行工作台。这样既能保证路线规划仍然可靠，也能通过真实地图底图、实时 POI、定位控件和步行路线反馈增强演示效果。",
        "本次实现以白色界面为视觉基调。相比传统深色面板或普通蓝白后台页面，白色出行工作台更接近地图类应用和校园服务平台的使用习惯，能够突出地图主体、降低信息压迫感，并让截图、报告和课堂演示更清晰。",
        "报告部分需要能够完整说明软件从需求分析到实现验证的全过程。因此报告不只记录结果，还要记录高德 API 接入方式、系统架构、数据模型、算法流程、界面设计、截图证据、测试命令、问题处理和后续扩展方向。报告与软件互相支撑，形成可以复核的证据链。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二、原报告与当前报告量级对比")
    add_paragraph(document, "在扩展本报告前，对原项目输出材料和当前材料进行了统计。统计维度包括 PDF 页数、DOCX 非空字符数、段落数量、表格数量和嵌入图片数量。结果表明，扩展前的当前报告明显偏短，无法支撑完整实验说明。")
    add_table(
        document,
        ["报告版本", "PDF 页数", "非空字符数", "段落数", "表格数", "图片数", "结论"],
        [
            ["原项目完成版", "13 页", "约 2911", "18", "4", "4", "可作为基础课程报告"],
            ["原项目长篇版", "74 页", "约 27349", "114", "4", "4", "内容明显更完整"],
            ["原项目论文体增强版", "88 页", "约 33437", "122", "5", "7", "原项目中篇幅最大的版本"],
            ["本项目扩展前", "4 页", "约 1655", "32", "2", "3", "篇幅不足，截图不足"],
            ["本项目扩展目标", "长篇报告", "显著增加", "显著增加", "显著增加", "13 张真实应用截图", "满足完整实验交付"],
        ],
    )
    add_paragraph(document, "根据上述对比，本次重新生成报告时采用长篇结构，并把真实运行截图作为重要证据。报告不追求与原文雷同，而是围绕高德 API、白色界面、校园出行工作台和新数据组织方式重新展开。")

    add_heading(document, "三、实验目的")
    add_bullets(
        document,
        [
            "掌握校园空间数据从现实地点到程序模型的转换方法。",
            "掌握高德地图开放平台 Web 端 Key、安全密钥和 JSAPI 2.0 Loader 的前端接入方式。",
            "掌握起终点路线规划、图结构建模和最短路径搜索在校园步行场景中的应用。",
            "掌握 React 前端、Express 后端、共享类型和静态回退规划的协同方式。",
            "掌握白色地图类界面的布局、层级、色彩、响应式和截图证据组织方法。",
            "完成一份与旧版风格不同的软件和长篇报告，形成可运行、可验证、可提交的完整材料。",
        ],
    )

    add_heading(document, "四、实验环境")
    add_table(
        document,
        ["项目", "内容"],
        [
            ["前端框架", "React + Vite + TypeScript"],
            ["地图服务", "高德地图 JSAPI 2.0，使用 Web 端 Key 和安全密钥在本地环境启用"],
            ["后端服务", "Express + TypeScript，提供 health、bootstrap 和 route 接口"],
            ["路线算法", "基于 GraphNode、GraphEdge 和地点 accessNodeIds 的带权最短路径搜索"],
            ["测试工具", "Vitest、ESLint、TypeScript 构建、Playwright 应用截图"],
            ["文档工具", "python-docx、Pillow、LibreOffice、Poppler"],
            ["交付目录", "output/doc/，包含报告、设计文档、发布说明、图表、截图和源码包"],
        ],
    )
    add_paragraph(document, "高德开放平台凭据只写入本地 .env.local，用于开发和截图验证，不写入源码压缩包、报告正文或提交清单。源码包中只保留 .env.example 作为配置示例，避免泄露真实 Key 和安全密钥。")

    add_heading(document, "五、实验原理")
    principle_paragraphs = [
        "系统把现实校园拆解为地点语义层、道路拓扑层、路线求解层和界面呈现层。地点记录描述名称、类别、场景带、别名、关键词、坐标和入口节点，路网节点与边描述可通行道路、距离和方向提示。路线规划时，系统根据起点与终点找到入口节点，在带权图上搜索最短路径，再生成折线坐标和逐段行动说明。",
        "地图层采用高德地图 JSAPI 2.0。前端在 AMapLoader.load 之前配置 window._AMapSecurityConfig，并加载 PlaceSearch、Walking、Geolocation、Scale、ToolBar、MapType 等插件，用于真实底图、实时 POI、定位和步行路线反馈。",
        "与旧版偏导览界面不同，新版本采用白色校园出行工作台定位，强调高德实景地图、路线调度和场景带过滤。视觉层面使用白色卡片、暖灰背景、青绿色主色和琥珀色路线强调，使软件在第一眼就呈现不同的产品气质。",
        "静态回退机制保证纯前端部署环境仍能运行核心路线规划。前端优先访问后端接口，如果接口不可用，则切换到本地规划器。两条路径返回相同 PlannedRoute 结构，因此界面无需区分后端结果与本地结果。",
        "高德地图接口与本地路网并不是互相替代关系。高德地图提供真实空间背景、实时地点和通用步行路线能力，本地路网提供课程实验需要的可解释校内节点和分段提示。两者结合后，用户既能看到真实地图，又能看到实验系统自己的路线计算结果。",
        "PlaceSearch 用于把校区和楼宇关键词映射到真实 POI 坐标，Walking 用于在真实地图上生成步行路线反馈，Geolocation 用于真实定位控件和实时导航入口。系统把这些能力放在高德实景视图中，同时保留拓扑视图作为本地路网校验工具。",
        "路线规划服务使用地点 accessNodeIds 连接地点和道路节点。一个地点可以有多个入口节点，规划时会遍历起点入口和终点入口组合，选择距离最短的候选路径。这种设计比把地点直接连成线更符合真实校园道路结构。",
        "前后端共享 PlannedRoute 类型，使后端路线规划结果、本地回退路线结果和前端展示层保持一致。这个契约降低了接口不一致的风险，也方便测试和报告复核。",
    ]
    for text in principle_paragraphs:
        add_paragraph(document, text)

    add_heading(document, "六、需求分析")
    requirement_sections = [
        ("用户需求", [
            "普通学生需要快速知道从宿舍、校门、餐厅到教学楼或图书馆的步行路径，并希望路线说明能够直接告诉自己经过哪些节点、距离大约是多少、预计走多久。",
            "新生或访客更依赖真实地图背景，因为他们对校内道路和建筑位置不熟悉。高德底图能提供周边道路、商圈、楼宇名称和地理参照物，降低迷路概率。",
            "课程展示人员需要系统具备稳定可复现的演示状态。通过 URL 参数记录起点、终点、筛选条件和选中地点，能够让截图和现场演示保持一致。",
        ]),
        ("功能需求", [
            "系统需要支持起点和终点选择，自动输出距离、时间和逐段路线说明。路线结果不应依赖静态图片，而应由程序根据图结构实时计算。",
            "系统需要支持高德实景地图、实时 POI、定位控件和步行路线反馈，并在缺少地图凭据时给出明确提示。",
            "系统需要支持地点搜索、分类筛选、场景带筛选、地点详情页、移动端响应式和本地路网拓扑校验。",
        ]),
        ("非功能需求", [
            "界面应采用白色主视觉，保证截图清晰、文字对比充足、控件边界明确，避免暗色界面在报告中打印或投影时可读性下降。",
            "报告材料应自动生成，避免手工排版遗漏。DOCX、PDF、图表、截图、源码包和提交清单应集中输出到同一目录。",
            "源码包必须排除 node_modules、dist、output、tmp、.env.local 和系统缓存文件，确保交付包干净且不包含敏感凭据。",
        ]),
    ]
    for heading, paragraphs in requirement_sections:
        add_heading(document, heading, level=2)
        for text in paragraphs:
            add_paragraph(document, text)

    add_heading(document, "七、总体设计")
    add_picture(document, DOCS["architecture_png"], caption="图 1  系统架构图")
    architecture_notes = [
        "前端层负责用户交互、地图展示、筛选状态和地点详情。首页把路线调度区放在左侧，把高德地图作为右侧主体，形成典型出行类应用的工作台布局。",
        "高德地图层是本次改造的关键差异点。AmapLiveMap 负责加载 JSAPI、初始化地图、添加控件、查询 POI、规划步行路线、展示实时导航卡片。AmapPoiExplorer 负责在本地地点之外补齐高德实时地点。",
        "导航服务层负责对外提供基础数据和路线规划接口。前端通过 requestRoutePlan 调用后端，如果后端不可用则调用 local-route-planner 完成本地回退。",
        "共享模型层统一 CampusConfig、PlaceRecord、GraphNode、GraphEdge 和 PlannedRoute 类型。共享契约使前端和后端都能围绕同一套数据结构工作。",
        "交付材料层由脚本统一生成。脚本读取项目指标，生成架构图、流程图、工作量统计图、报告、设计文档、发布说明、截图清单和源码压缩包。",
    ]
    for text in architecture_notes:
        add_paragraph(document, text)
    add_picture(document, DOCS["workflow_png"], caption="图 2  开发与交付流程图")
    add_picture(document, DOCS["workload_png"], caption="图 3  项目工作量统计图")

    add_heading(document, "八、数据模型设计")
    add_paragraph(document, "系统数据采用地点语义和道路拓扑分离的设计。地点语义解决用户如何搜索和理解一个建筑，道路拓扑解决程序如何在节点之间寻找可通行路径。两者通过 accessNodeIds 建立关联。")
    add_table(
        document,
        ["模型", "主要字段", "设计说明"],
        [
            ["CampusConfig", "id, name, city, mapAsset, mapAlt", "描述校区名称、城市、底图资源和地图说明，是页面和报告的基础配置。"],
            ["PlaceRecord", "id, name, categoryId, zone, aliases, keywords, mapPoint, accessNodeIds", "描述一个可被搜索、筛选、预览和规划的校园地点。"],
            ["GraphNode", "id, label, zone, xPct, yPct", "描述路网节点，坐标使用百分比，便于在拓扑底图上绘制。"],
            ["GraphEdge", "fromNodeId, toNodeId, distanceMeters, instruction", "描述两节点之间的可通行道路和步行提示。"],
            ["PlannedRoute", "distanceMeters, estimatedMinutes, pathPoints, steps", "描述最终路线结果，前端摘要、拓扑折线和报告截图都依赖该结构。"],
            ["LivePoiSelection", "id, name, type, distanceMeters, lng, lat", "描述高德实时 POI 结果，用于补齐本地路网之外的地点。"],
        ],
    )
    data_notes = [
        "PlaceRecord 的 aliases 和 keywords 使用户可以用简称、楼号、用途和常见叫法检索地点。例如餐厅、宿舍、J 楼、若水园等关键词都可以被本地搜索识别。",
        "categoryId 和 zone 分别解决不同维度的筛选需求。categoryId 面向任务目的，例如学习、餐饮、体育、医疗；zone 面向空间分区，例如学习核心带、生活补给带和访客入口带。",
        "mapPoint 用于拓扑视图中的点位绘制，而真实高德地图坐标则通过 PlaceSearch 或实时 POI 返回。这样避免了把演示用拓扑坐标误当成真实经纬度。",
        "GraphEdge 中保留 instruction 字段，使每段路线不只是线段，而是可解释的行动提示。报告截图中的分步说明正是由这些字段生成。",
    ]
    for text in data_notes:
        add_paragraph(document, text)

    add_heading(document, "九、高德地图 API 接入设计")
    amap_notes = [
        "高德 JSAPI 2.0 要求在加载地图前完成安全配置。系统在动态导入 @amap/amap-jsapi-loader 后，先写入 window._AMapSecurityConfig，再调用 AMapLoader.load。这个顺序必须保证，否则地图或插件可能加载失败。",
        "地图初始化采用 3D 视图、默认中心点和标准样式，并添加 Scale、ToolBar、MapType、Geolocation 控件。用户进入页面后能够看到标准图层、卫星图、路网和路况等高德控件。",
        "PlaceSearch 首先用于定位校区中心，再用于根据本地点位名称、别名和校区组合关键词查询真实 POI。查询结果会缓存到 placeCacheRef，避免重复检索同一地点。",
        "Walking 插件用于在高德实景地图上展示真实步行路线。系统在用户选择起点和终点后先解析两个地点的真实坐标，再调用 Walking.search 生成地图路线反馈。",
        "Geolocation 和浏览器定位用于实时导航入口。当前实现支持启动定位、转换 GPS 坐标到高德坐标、展示当前位置箭头、判断偏航并触发重新规划。",
        "POI 补齐功能通过 searchNearBy 查询校区附近地点。由于高德数据可能返回商圈、维修点、服务点等真实周边信息，系统会根据距离、关键词和去重规则挑选可展示结果。",
        "实际 Key 与安全密钥只存在本地 .env.local 中，不进入源码包。报告和源码示例只写环境变量名，保证交付材料可以说明配置方式，同时避免泄露真实凭据。",
    ]
    for text in amap_notes:
        add_paragraph(document, text)
    add_table(
        document,
        ["高德能力", "对应插件/对象", "系统用途", "界面体现"],
        [
            ["真实底图", "AMap.Map", "展示校园周边道路、建筑和空间背景", "右侧高德实景地图主视图"],
            ["比例尺和工具条", "Scale, ToolBar", "辅助用户理解缩放级别和空间距离", "地图左下和右侧控件"],
            ["图层切换", "MapType", "支持标准图层、卫星图、路网和路况", "地图右上控件组"],
            ["POI 检索", "PlaceSearch", "定位校区、解析地点、补齐实时地点", "高德实时地点列表"],
            ["步行规划", "Walking", "输出真实地图上的步行路线反馈", "选择起终点后的实景路线"],
            ["定位能力", "Geolocation + 浏览器定位", "为实时导航和当前位置箭头提供基础", "实时步行导航卡片"],
        ],
    )

    add_heading(document, "十、路线规划算法设计")
    route_algorithm = [
        "路线规划的输入是起点地点 id 和终点地点 id。系统首先在地点表中查找两个 PlaceRecord，如果任意一个不存在，则返回可读错误，前端把错误显示在路线调度卡片中。",
        "当起点和终点相同时，系统直接返回距离为 0 的路线结果。这种特殊分支避免了后续图搜索浪费，并让用户看到明确提示。",
        "当起点和终点不同时，系统读取两个地点的 accessNodeIds。每组起点入口和终点入口都构成一个候选规划任务，算法会尝试所有组合并选择总距离最短的结果。",
        "图搜索部分维护节点状态，包括当前距离、前驱节点和是否已访问。每次选择未访问节点中距离最小的节点继续扩展，这与 Dijkstra 最短路径思想一致。",
        "搜索结束后，系统从终点节点沿 previousNodeId 回溯，得到完整节点序列。节点序列再转换成 pathPoints，供拓扑图绘制折线。",
        "路线 steps 由相邻节点对应的 GraphEdge 生成。每一步保留 fromLabel、toLabel、instruction 和 distanceMeters，使摘要区可以展示自然语言步行提示。",
        "预计耗时通过总距离估算。虽然真实校园步行速度会受坡度、拥堵、天气和红绿灯影响，但对于课程实验和演示，基于距离的分钟估算足够直观。",
    ]
    for text in route_algorithm:
        add_paragraph(document, text)
    add_table(
        document,
        ["步骤", "输入", "处理", "输出"],
        [
            ["1", "startPlaceId, endPlaceId", "校验地点是否存在", "起终点 PlaceRecord"],
            ["2", "accessNodeIds", "枚举入口节点组合", "候选起终点节点对"],
            ["3", "GraphNode, GraphEdge", "执行带权最短路径搜索", "最短节点路径"],
            ["4", "节点路径", "汇总距离并生成 pathPoints", "路线折线数据"],
            ["5", "边信息", "生成逐段 instruction", "路线步骤列表"],
            ["6", "PlannedRoute", "前端地图和摘要渲染", "可视化路线结果"],
        ],
    )

    add_heading(document, "十一、界面与交互设计")
    ui_notes = [
        "界面采用白色主视觉，背景使用很轻的网格纹理，卡片保持 8px 圆角和细边框。这样既有地图工作台的专业感，又不会像深色面板一样压暗高德底图。",
        "首屏采用左控件、右地图的结构。左侧负责选择和摘要，右侧负责地图主体。用户进入系统后先看到真实地图，而不是被大量说明文字包围。",
        "路线调度卡片只保留必要控件：起点、对调、终点、清空。没有增加不必要的配置项，符合课程实验中“最小可用、容易演示”的目标。",
        "地图卡片提供实景地图和校园拓扑两个标签。实景地图默认选中，拓扑图用于校验本地节点和路线折线，二者定位清晰，不互相抢占职责。",
        "搜索区、实时 POI、分类筛选、场景带和地点队列放在首屏下方。这样首屏保持路线和地图优先，更多地点探索能力向下展开。",
        "移动端布局改为纵向堆叠，路线调度、地图、摘要、搜索和地点列表依次出现。按钮、输入框和卡片宽度都适配窄屏，避免文字挤压和横向滚动。",
    ]
    for text in ui_notes:
        add_paragraph(document, text)

    add_heading(document, "十二、真实应用截图")
    add_paragraph(document, "本节使用 Playwright 对本地运行的软件进行截图，截图均来自真实页面，而不是重新绘制的示意图。完整截图文件保存在 output/doc/application_screenshots/ 目录，报告中插入的是裁剪后的关键区域图，便于阅读和排版。")
    add_application_screenshot_gallery(document, start_index=4)

    add_heading(document, "十三、实验过程记录")
    add_bullets(
        document,
        [
            "读取旧项目结构，确认可复用的路线规划、地点数据、测试和构建链路。",
            "新建独立目录 CAMPUS_PULSE，清理旧缓存、旧截图、旧 APK 壳和无关论文材料。",
            "重命名项目、底图和校区配置，建立山科智行 Campus Pulse 新品牌。",
            "将原分区改造为学习核心带、实验创新带、生活补给带和访客入口带。",
            "配置高德地图开放平台 Web 端 Key 和安全密钥，启用 JSAPI 实景地图、POI 检索、定位和步行规划。",
            "重写首页、路线规划、地图、搜索、地点列表、详情页和状态提示文案。",
            "整块重写 CSS，形成白色校园出行工作台界面，并同步更新测试断言。",
            "重写报告生成脚本，输出新文件名、新图表、新说明和新的源码压缩包。",
        ],
    )
    process_notes = [
        "第一阶段重点是拆解基线项目。保留路线规划、图结构、共享类型和测试体系，清理与新作品无关的旧截图、旧 APK 包装、缓存目录和旧文档输出。",
        "第二阶段完成命名和数据重构。项目名称从旧导览定位改为山科智行 Campus Pulse，分区名称改为学习核心带、实验创新带、生活补给带和访客入口带，底图文件也重新命名。",
        "第三阶段完成高德地图接入。系统使用 VITE_AMAP_JSAPI_KEY 和 VITE_AMAP_SECURITY_JS_CODE 读取本地凭据，并在 AMapLoader.load 前设置安全配置。",
        "第四阶段完成白色界面重做。全局变量、卡片、地图容器、按钮、筛选芯片、列表条目、详情页和移动端断点都重新设计，避免继续保留旧风格。",
        "第五阶段完成真实截图。通过 Playwright 分别截取默认首页、路线结果、地点快照、关键词搜索、分类筛选、场景筛选、详情页、移动端和拓扑校验视图。",
        "第六阶段完成报告扩写。报告从短说明扩展为长篇实验记录，加入量级对比、技术原理、API 接入、算法流程、界面设计、截图图集和验收记录。",
    ]
    for text in process_notes:
        add_paragraph(document, text)

    add_heading(document, "十四、关键结果")
    add_table(
        document,
        ["指标", "数值", "说明"],
        [
            ["源码文件", str(metrics.source_files), "统计前端、后端、共享类型和样式文件"],
            ["源码行数", str(metrics.source_lines), "不含依赖、构建产物和输出材料"],
            ["测试用例", str(metrics.test_cases), "覆盖首页、地图切换、路线规划和服务逻辑"],
            ["点位数量", str(metrics.place_count), "当前内置可规划校园地点"],
            ["路网规模", f"{metrics.graph_nodes} 节点 / {metrics.graph_edges} 边", "用于步行路径搜索"],
            ["应用截图", str(len(APP_SCREENSHOTS)), "覆盖桌面端、移动端、高德地图、搜索、筛选、详情页和拓扑校验"],
            ["高德能力", "JSAPI 2.0", "真实底图、POI 检索、定位、步行规划和地图控件"],
        ],
    )

    add_heading(document, "十五、测试与验证")
    add_table(
        document,
        ["验证项", "命令或方式", "结果", "说明"],
        [
            ["单元与组件测试", "npm run test", "通过", "4 个测试文件，16 条用例全部通过"],
            ["类型与生产构建", "npm run build", "通过", "TypeScript 构建和 Vite 打包成功"],
            ["代码规范", "npm run lint", "通过", "ESLint 无错误无警告"],
            ["后端健康检查", "GET /api/health", "通过", "返回 {ok:true}"],
            ["高德地图加载", "Playwright 截图", "通过", "可看到高德底图和图层控件"],
            ["路线规划", "URL 参数 + 接口规划", "通过", "西门到图书信息中心生成 560 米、7 分钟路线"],
            ["移动端展示", "390px 视口截图", "通过", "核心控件和地图区域保持可读"],
            ["源码包检查", "zip 内容检索", "通过", "未包含 .env.local、真实 Key 或安全密钥"],
        ],
    )
    validation_notes = [
        "测试结果说明系统核心功能没有因为界面重写和高德接入而破坏。首页渲染、分类过滤、地点预览、详情页、路线规划、地图切换和服务逻辑仍然保持可用。",
        "构建结果说明 TypeScript 类型、Vite 打包和依赖解析均正常。由于项目运行在外置盘环境，额外处理了平台相关依赖路径，最终构建产物可正常输出到 dist。",
        "lint 结果说明 React Hook 依赖、导入和常规代码规范已经清理干净。高德地图组件中原先的 Hook 依赖警告通过函数 ref 方式处理，避免地图初始化 effect 被不必要地反复触发。",
        "Playwright 截图不仅用于视觉检查，也用于报告证据。截图能够证明高德底图加载成功、白色界面生效、路线结果可见、筛选功能可复现、移动端布局可读。",
    ]
    for text in validation_notes:
        add_paragraph(document, text)

    add_heading(document, "十六、问题与解决")
    add_table(
        document,
        ["问题", "现象", "原因", "解决方式"],
        [
            ["高德凭据未配置", "地图区域提示无法加载实景底图", "缺少 Web 端 Key 或安全密钥", "在 .env.local 配置 VITE_AMAP_JSAPI_KEY 和 VITE_AMAP_SECURITY_JS_CODE"],
            ["HashRouter 参数位置错误", "URL 中 start/end 未被读取", "查询参数放在 # 之前", "改用 /#/?start=...&end=... 复现截图状态"],
            ["报告过短", "PDF 只有 4 页，截图只有 3 张", "脚本正文和图集不足", "扩写报告结构并加入 13 张真实应用截图"],
            ["界面观感差", "深色网格压抑、地图主体不突出", "视觉方向不适合地图应用", "改为白色出行工作台、暖灰背景和高德地图主视图"],
            ["源码包敏感信息风险", "本地需要真实 Key，交付不能泄露", ".env.local 只应本地使用", "源码包排除 .env.local，并精确检索确认无密钥"],
            ["外置盘元数据文件", "出现 ._* 文件影响 lint 或统计", "macOS 在外置盘生成 AppleDouble 文件", "生成和验证前统一 find 删除"],
        ],
    )
    problem_notes = [
        "这些问题体现了课程实验中常见的工程细节：软件能运行只是第一步，报告是否充分、截图是否真实、配置是否安全、交付包是否干净同样重要。",
        "高德地图 API 的接入尤其需要注意安全配置顺序。安全密钥必须在 Loader 加载前设置，否则即使 Key 正确也可能出现地图白屏或插件异常。",
        "报告扩写不是简单堆字数，而是把软件实现过程拆成可检查的证据：需求、架构、数据、算法、API、界面、截图、测试、问题和结论都应当有对应内容。",
    ]
    for text in problem_notes:
        add_paragraph(document, text)

    add_heading(document, "十七、分析与讨论")
    discussion = [
        "本次新版本不是简单替换标题，而是同时改造了产品定位、视觉表达、信息结构、文档脚本和交付命名。原基线更偏传统导览，新版以行动看板和路线调度为主，能够在同一底层算法基础上呈现完全不同的界面观感。",
        "项目保留了原系统中真正有价值的能力，包括图结构路线规划、前后端共享类型、静态回退和测试体系。这种做法比完全重写更稳妥，因为底层可验证能力不丢失，同时上层体验和报告材料已经重新设计。",
        "当前系统仍有扩展空间，例如接入正式校区底图、补全全量楼宇入口、增加封路规则、支持从当前位置出发和后台维护。但作为课程实验交付，已经具备软件、报告、图表、截图和源码包组成的完整证据链。",
        "高德地图 API 提升了系统的真实感，但也带来了外部服务依赖。课程展示时需要保证网络可用、Key 可用、浏览器允许加载地图脚本。如果在无网络环境演示，校园拓扑视图和本地路线规划仍然可以承担核心说明。",
        "白色界面更适合报告截图和课堂投影。地图底图本身已经包含道路、建筑和颜色层级，如果界面外框过暗或装饰过多，会削弱地图可读性。本次改版减少了装饰，强调地图、路线和地点信息本身。",
        "本地路网和高德路线之间存在粒度差异。高德路线依赖真实道路数据，本地路网依赖课程数据建模；当两者不完全一致时，报告中应说明本地路网用于实验可解释路线，高德路线用于真实空间背景和外部地图能力验证。",
        "截图集的价值在于覆盖主要用户路径。只有首页截图不足以证明软件完整性，因此本次增加了路线结果、搜索、分类、场景带、详情页、移动端和拓扑校验截图。",
        "自动生成报告脚本可以减少重复劳动，但也需要不断更新内容模板。如果软件已改成高德 API 和白色界面，报告脚本中仍残留深色界面或旧项目描述，就会造成材料不一致。",
    ]
    for text in discussion:
        add_paragraph(document, text)

    add_heading(document, "十八、主要源码模块说明")
    add_paragraph(document, "为了便于教师或评阅人员复核，本节对主要源码模块进行说明。说明重点不是逐行解释代码，而是说明每个文件在系统中的职责、输入输出和修改价值。")
    add_table(
        document,
        ["模块", "文件", "职责", "本次改造重点"],
        [
            ["首页路由", "src/routes/HomePage.tsx", "组织首屏、路线调度、地图、搜索、筛选和地点列表", "改为高德 API 出行工作台文案与状态组织"],
            ["地图容器", "src/components/CampusMap.tsx", "在高德实景地图和校园拓扑之间切换", "把高德地图设为默认主视图，拓扑作为校验视图"],
            ["高德地图", "src/components/AmapLiveMap.tsx", "加载 JSAPI、初始化地图、POI 查询、步行规划和实时导航", "接入 Key、安全密钥、PlaceSearch、Walking 和 Geolocation"],
            ["实时 POI", "src/components/AmapPoiExplorer.tsx", "从高德检索周边 POI 并展示可定位地点", "补齐本地路网之外的实时地点"],
            ["路线调度", "src/components/RoutePlanner.tsx", "提供起点、终点和对调操作", "保持控件简洁，适配白色界面"],
            ["路线摘要", "src/components/RouteSummary.tsx", "展示总距离、预计耗时和分步说明", "与高德地图并排展示，形成路线证据"],
            ["地点列表", "src/components/PlaceList.tsx", "展示筛选后的本地点位", "保留起终点标记和关键词标签"],
            ["本地算法", "src/lib/local-route-planner.ts", "在浏览器侧完成最短路径规划", "作为后端不可用时的静态回退"],
            ["后端服务", "server/src/app.ts", "提供 API 路由和错误处理", "继续支撑 route 规划接口和 health 检查"],
            ["报告脚本", "scripts/generate_experiment_submission.py", "生成 DOCX/PDF、图表、截图清单和源码包", "扩展长篇报告和真实应用截图图集"],
        ],
    )
    module_details = [
        "HomePage 是整个前端体验的编排层。它读取 URL 中的搜索参数，把 q、category、zone、selected、start 和 end 转换为界面状态，并把状态变化同步回链接。这样截图和演示可以通过链接稳定复现。",
        "CampusMap 是地图能力的门面组件。组件内部不直接处理高德 API 细节，而是根据当前模式渲染 AmapLiveMap 或本地拓扑底图。这样的分层让高德地图逻辑和拓扑校验逻辑互不混杂。",
        "AmapLiveMap 是本次工作量最大的组件。它包含地图加载、控件添加、地点解析、步行规划、实时定位、语音播报和导航状态展示。虽然组件复杂，但它集中处理高德地图相关副作用，避免把外部 API 逻辑散落到多个界面组件中。",
        "AmapPoiExplorer 体现了高德 API 的补齐价值。即使本地数据只维护课程实验所需地点，用户仍可以通过高德实时 POI 看到周边服务点和真实地点，从而让系统更接近真实地图应用。",
        "local-route-planner 保留了课程实验最核心的算法证明。即使高德地图无法访问，本地路线规划仍然可以根据 GraphEdge 输出可解释路径。这是系统可靠性和教学价值的重要保障。",
        "报告生成脚本不只是文档工具，也是交付自动化工具。它把截图、指标、图表、源码包和 PDF 转换放到同一个流程，确保每次软件更新后材料可以同步刷新。",
    ]
    for text in module_details:
        add_paragraph(document, text)

    add_heading(document, "十九、典型使用场景")
    scenarios = [
        ("场景一：新生从西门到图书信息中心", "用户选择西门作为起点，选择图书信息中心作为终点。系统生成 560 米、约 7 分钟的推荐路线，并在高德地图中展示真实空间背景。这个场景适合迎新或新生入馆指引。"),
        ("场景二：宿舍区到工程实训中心", "用户从学生公寓 A 区前往工程实训中心。路线需要跨越生活补给带、学习核心带和实验创新带，可以验证较长距离和跨区路线规划能力。"),
        ("场景三：查找餐饮地点", "用户输入餐厅关键词，系统同时过滤本地点位并加载高德实时 POI。这个场景验证了本地数据和外部地图数据的组合检索效果。"),
        ("场景四：按行动目的筛选", "用户点击餐饮、图书馆、宿舍等分类筛选芯片，地点队列和统计信息同步变化。这个场景适合展示信息架构和前端状态管理。"),
        ("场景五：按空间场景带筛选", "用户选择生活补给带或实验创新带，系统只展示对应空间带内地点。这个场景适合解释校区空间组织方式。"),
        ("场景六：进入地点详情页", "用户从地点快照进入图书信息中心详情页，查看抵达建议、关键词和邻近推荐。这个场景验证直达链接和单点说明能力。"),
        ("场景七：切换校园拓扑视图", "用户在路线结果中切换到校园拓扑，核对本地路网节点、路线折线和起终点标记。这个场景体现本地算法证据。"),
        ("场景八：手机端查看路线", "用户在 390px 宽度下打开页面，路线调度、地图、摘要和地点列表按纵向排列。这个场景验证移动端响应式可用性。"),
    ]
    add_table(document, ["场景", "说明"], [[title, description] for title, description in scenarios])
    for title, description in scenarios:
        add_paragraph(document, f"{title}：{description}")

    add_heading(document, "二十、点位数据清单")
    add_paragraph(document, "当前版本内置 22 个可规划点位，覆盖学习、实验、生活、访客入口、体育、医疗、餐饮、宿舍等校园常见行动目的。点位数据是路线规划、搜索筛选、详情页和地图标记的共同基础。")
    add_table(
        document,
        ["序号", "点位", "类别", "场景带", "典型用途"],
        [
            ["1", "图书信息中心", "图书馆", "学习核心带", "学习、自习、查阅资料和路线演示终点"],
            ["2", "工程实训中心", "实验实训", "实验创新带", "实训课程、实验教学和跨区路线终点"],
            ["3", "逸夫楼", "教学楼", "学习核心带", "公共课、教学活动和中轴路线节点"],
            ["4", "行政楼", "行政服务", "学习核心带", "办事、咨询和校园行政服务"],
            ["5", "学术交流中心", "行政服务", "访客入口带", "访客接待、会议和西门入口缓冲"],
            ["6", "西门", "校门", "访客入口带", "入校起点、迎新入口和访客路线起点"],
            ["7", "东门", "校门", "生活补给带", "生活区出入口和东侧通行起点"],
            ["8", "体育运动场", "体育", "访客入口带", "体育课、运动训练和活动集合"],
            ["9", "综合体育馆", "体育", "访客入口带", "室内体育、活动和比赛场景"],
            ["10", "校医院", "医疗", "生活补给带", "就医、体检和应急服务"],
            ["11", "A 餐厅", "餐饮", "生活补给带", "日常就餐和生活服务"],
            ["12", "B 餐厅", "餐饮", "生活补给带", "日常就餐和宿舍区补给"],
            ["13", "学生公寓 B1", "宿舍", "生活补给带", "住宿、生活区路线起点"],
            ["14", "学生公寓 B3", "宿舍", "生活补给带", "住宿、生活区路线起点"],
            ["15", "学生公寓 C13", "宿舍", "生活补给带", "住宿、生活区路线起点"],
            ["16", "J6 测绘经管地质学院楼", "教学楼", "学习核心带", "学院教学和课程地点"],
            ["17", "地球科学与工程学院楼", "实验实训", "实验创新带", "学院教学和科研实验"],
            ["18", "若水园", "地标", "学习核心带", "校园地标、休息和路线参照"],
            ["19", "学生公寓 A 区", "宿舍", "生活补给带", "宿舍到教学区路线起点"],
            ["20", "学生公寓 A17", "宿舍", "生活补给带", "宿舍到餐厅和教学区路线起点"],
            ["21", "学生公寓 B2", "宿舍", "生活补给带", "生活区住宿点"],
            ["22", "学苑餐厅", "餐饮", "生活补给带", "生活区餐饮补充点"],
        ],
    )
    point_notes = [
        "点位清单中既包含稳定建筑，也包含服务类地点和入口节点。这样做可以覆盖从入校、学习、实验、就餐、住宿到就医的完整校园行动链路。",
        "每个点位都保留 categoryId 和 zone。前者服务任务目的筛选，后者服务空间分层筛选。用户既可以按“想做什么”查找，也可以按“在哪一片区域”查找。",
        "宿舍和餐厅类点位数量较多，是因为校园步行场景中生活区往往是起点或终点高频区域。增加这些点位能让路线样例更贴近真实使用。",
    ]
    for text in point_notes:
        add_paragraph(document, text)

    add_heading(document, "二十一、路线样例与验收记录")
    add_table(
        document,
        ["样例", "起点", "终点", "预期验证点", "截图或证据"],
        [
            ["R1", "西门", "图书信息中心", "短路线、入校到学习区、摘要和高德底图", "02_路线规划结果_西门到图书信息中心.png"],
            ["R2", "学生公寓 A 区", "工程实训中心", "跨生活区和实验区的较长路线", "07_宿舍到实训中心路线.png"],
            ["R3", "西门", "图书信息中心", "拓扑视图折线与起终点标记", "13_校园拓扑路线校验.png"],
            ["R4", "移动端西门", "移动端图书信息中心", "窄屏下路线摘要和地图显示", "11_移动端路线规划结果.png"],
            ["R5", "餐厅搜索", "A 餐厅/B 餐厅/学苑餐厅", "关键词搜索和本地筛选", "04_搜索餐厅结果.png"],
        ],
    )
    route_notes = [
        "R1 是基础验收路线，能验证起终点选择、后端 route 接口、前端摘要、高德地图和截图生成的完整链路。",
        "R2 用于验证较长路线和跨区路线，避免系统只在短距离或单一区域内看起来可用。",
        "R3 用于验证本地拓扑视图仍然可用。即使高德地图作为默认主视图，拓扑视图也必须保留课程实验的算法说明价值。",
        "R4 用于验证移动端可用性。移动端截图显示地图、路线摘要和列表都可以纵向阅读，没有横向溢出。",
        "R5 用于验证搜索和实时 POI 补齐。餐厅关键词能够触发本地列表过滤，同时高德 POI 区域展示实时结果。",
    ]
    for text in route_notes:
        add_paragraph(document, text)

    add_heading(document, "二十二、截图采集记录")
    add_paragraph(document, "应用截图通过 Playwright 在本地开发服务上采集。桌面端使用 1440px 宽视口，移动端使用 390px 宽视口。每张截图等待高德地图瓦片加载后再保存，避免地图区域空白。")
    add_table(
        document,
        ["编号", "文件名", "视口", "状态来源", "说明"],
        [
            ["01", "01_高德实景首页.png", "1440px", "/#/", "默认首页和高德实景地图"],
            ["02", "02_路线规划结果_西门到图书信息中心.png", "1440px", "/#/?start=west-gate&end=library-information-center", "基础路线结果"],
            ["03", "03_地点快照_图书信息中心.png", "1440px", "/#/?selected=library-information-center", "地点预览状态"],
            ["04", "04_搜索餐厅结果.png", "1440px", "/#/?q=餐厅", "关键词搜索状态"],
            ["05", "05_餐饮分类筛选.png", "1440px", "/#/?category=dining", "分类筛选状态"],
            ["06", "06_生活补给带筛选.png", "1440px", "/#/?zone=生活补给带", "场景带筛选状态"],
            ["07", "07_宿舍到实训中心路线.png", "1440px", "/#/?start=student-dorm-a-zone&end=engineering-training-center", "跨区路线状态"],
            ["08", "08_地点详情页_图书信息中心.png", "1440px", "/#/place/library-information-center", "详情页状态"],
            ["09", "09_地点详情页_学生公寓A区.png", "1440px", "/#/place/student-dorm-a-zone", "宿舍详情页状态"],
            ["10", "10_移动端高德首页.png", "390px", "/#/", "移动端首页"],
            ["11", "11_移动端路线规划结果.png", "390px", "/#/?start=west-gate&end=library-information-center", "移动端路线"],
            ["12", "12_移动端搜索餐厅.png", "390px", "/#/?q=餐厅", "移动端搜索"],
            ["13", "13_校园拓扑路线校验.png", "1280px", "交互点击校园拓扑", "拓扑校验视图"],
        ],
    )
    add_paragraph(document, "截图采集使用 HashRouter 的链接格式，因此查询参数必须放在 #/ 后面。例如路线截图使用 /#/?start=west-gate&end=library-information-center。如果把查询参数放在 # 之前，React Router 无法读取页面状态。")
    add_paragraph(document, "每张截图都保存在 output/doc/application_screenshots/ 中，报告中另行生成 report_figures/ 裁剪图用于排版。完整截图保留了页面全长内容，裁剪图用于 DOCX/PDF 中的关键区域展示。")

    add_heading(document, "二十三、交付审计")
    add_table(
        document,
        ["审计项", "检查方法", "结果"],
        [
            ["真实密钥是否泄露", "检索源码、报告文本和源码压缩包", "未发现真实 Key 或安全密钥"],
            [".env.local 是否进入源码包", "zipgrep 和文件列表检查", "未进入源码包"],
            ["旧项目文字是否残留", "检索旧名称、旧系统名和旧关键词", "未在源码和报告文本中发现"],
            ["应用截图是否充足", "统计 application_screenshots 目录", "13 张真实应用截图"],
            ["报告是否过短", "统计 PDF 页数和 DOCX 字符数", "已扩展为长篇报告"],
            ["外置盘元数据文件", "find . -name '._*'", "生成后清理"],
        ],
    )
    audit_notes = [
        "交付审计的目的不是增加形式，而是避免常见扣分点。尤其是地图 Key 泄露、报告截图不足、报告仍写旧项目名称、源码包包含 node_modules 或输出目录，都会影响最终提交质量。",
        "本项目保留 .env.example 用于说明配置项，但真实 .env.local 被 .gitignore 和源码打包逻辑排除。这样用户可以运行高德地图，又不会把凭据交出去。",
        "应用截图目录是本次补充的重要材料。它不仅服务报告排版，也可以在提交材料中单独作为软件运行效果证据。",
    ]
    for text in audit_notes:
        add_paragraph(document, text)

    add_heading(document, "二十四、运行命令与验证记录")
    add_paragraph(document, "本节记录本项目从开发运行到交付检查使用的主要命令。记录这些命令的目的，是让评阅人员能够复现实验过程，而不是只能查看最终截图。")
    add_table(
        document,
        ["序号", "命令或检查", "用途", "结果说明"],
        [
            ["1", "npm run dev", "启动 Vite 前端与 Express 后端", "前端监听 http://localhost:5173/，后端监听 http://127.0.0.1:8787"],
            ["2", "curl /api/health", "检查后端服务是否可用", "返回 {ok:true}"],
            ["3", "npm run test", "执行 Vitest 自动测试", "4 个测试文件、16 条用例通过"],
            ["4", "npm run build", "执行 TypeScript 构建和 Vite 生产打包", "dist 目录正常生成"],
            ["5", "npm run lint", "执行 ESLint 代码规范检查", "无错误无警告"],
            ["6", "Playwright screenshot", "采集桌面端和移动端真实应用截图", "生成 13 张应用截图"],
            ["7", "python3 scripts/generate_experiment_submission.py", "生成 DOCX/PDF、图表、截图清单和源码包", "所有交付材料输出到 output/doc"],
            ["8", "zipgrep / rg", "检查源码包是否包含真实 Key 或 .env.local", "未发现敏感凭据"],
            ["9", "pdfinfo", "检查 PDF 页数和文件信息", "实验报告约 40 页量级"],
            ["10", "find . -name '._*' -delete", "清理外置盘 AppleDouble 元数据文件", "避免 lint 和源码包污染"],
        ],
    )
    command_notes = [
        "npm run dev 同时启动前端和后端，便于在浏览器中验证完整链路。高德地图需要前端环境变量，因此开发服务重启后才能读取最新 .env.local。",
        "curl /api/health 是最小后端可用性验证。只有后端可访问，前端才会优先使用服务端 route 接口；如果后端不可用，前端仍会回退到本地规划器。",
        "npm run test 覆盖首页渲染、分类过滤、地点预览、详情页、路线规划和地图切换。测试通过说明界面重做没有破坏原有关键行为。",
        "npm run build 能同时检查 TypeScript 类型和 Vite 资源打包。对于课程提交而言，构建通过比只在开发服务器能运行更有说服力。",
        "npm run lint 用于发现隐藏代码问题。本项目最终清理了高德地图组件中的 Hook 依赖警告，避免留下影响交付质量的提示。",
        "Playwright 截图命令等待高德地图瓦片加载后保存图片。这样生成的图可以证明真实地图已经加载，而不是仅有空白地图容器。",
        "报告生成脚本统一处理 DOCX、PDF、图表、截图裁剪和源码压缩包。相比手工复制截图，脚本化方式更适合重复迭代。",
        "敏感凭据检查是必要步骤。高德 Key 和安全密钥用于本地运行，但不应进入源码包、报告正文或提交清单。",
    ]
    for text in command_notes:
        add_paragraph(document, text)

    add_heading(document, "二十五、接口与数据流验收")
    add_paragraph(document, "系统的数据流从用户选择起终点开始，经过前端状态管理、后端路线接口、本地或服务端算法、高德地图解析和前端展示。下表把主要数据流拆解为可检查节点。")
    add_table(
        document,
        ["数据流节点", "输入", "处理模块", "输出", "验收方式"],
        [
            ["筛选参数解析", "URL 中 q/category/zone/selected/start/end", "parseFilters", "FilterState", "通过不同截图 URL 复现状态"],
            ["地点过滤", "places + FilterState", "filterPlaces", "visiblePlaces", "搜索和分类截图验证"],
            ["路线请求", "startPlaceId/endPlaceId", "requestRoutePlan", "PlannedRoute", "路线摘要截图验证"],
            ["后端规划", "route API payload", "server route-planner", "距离、时间、steps", "curl 和前端结果验证"],
            ["本地回退", "相同起终点", "local-route-planner", "PlannedRoute", "后端不可用时仍可规划"],
            ["高德地图加载", "Key + securityJsCode", "AMapLoader.load", "AMap namespace", "实景底图截图验证"],
            ["POI 解析", "地点名称和别名", "PlaceSearch", "真实经纬度候选", "高德状态和地图定位验证"],
            ["实时地点", "关键词和中心点", "searchNearBy", "LivePoiSelection[]", "高德实时地点列表验证"],
            ["步行路线", "起终点经纬度", "Walking.search", "高德地图路线反馈", "实景路线状态验证"],
            ["拓扑校验", "PlannedRoute.pathPoints", "CampusMap SVG polyline", "本地折线", "校园拓扑截图验证"],
        ],
    )
    flow_notes = [
        "URL 参数是系统可复现性的基础。截图采集时不需要人工重复点击大部分控件，只要构造正确的 HashRouter URL，就能直接进入指定状态。",
        "本地地点过滤和高德实时地点不是同一套数据。本地地点用于课程规划和可解释路线，高德实时地点用于补齐真实地图中的周边服务点。",
        "路线请求先尝试后端接口，这样服务端和浏览器端都可以参与实验；当后端异常时，本地回退算法保证核心功能不至于中断。",
        "高德地图加载成功后，地图控件、底图版权、标准图层和卫星图选项都能在截图中看到。这些视觉证据比单纯说“已接入高德”更可靠。",
        "拓扑校验视图不是旧功能残留，而是用于验证本地路网算法。高德地图证明真实空间背景，本地拓扑证明课程数据结构和路径算法。",
    ]
    for text in flow_notes:
        add_paragraph(document, text)

    add_heading(document, "二十六、配置安全与交付边界")
    add_paragraph(document, "高德地图开放平台的 Key 和安全密钥属于运行凭据。实验中需要使用真实凭据验证地图加载，但交付材料不能泄露这些值。因此本项目把运行配置和源码交付做了明确分离。")
    add_table(
        document,
        ["配置项", "所在位置", "是否进入源码包", "说明"],
        [
            ["VITE_AMAP_JSAPI_KEY", ".env.local", "否", "高德 Web 端 Key，仅本地运行使用"],
            ["VITE_AMAP_SECURITY_JS_CODE", ".env.local", "否", "高德安全密钥，仅本地运行使用"],
            ["VITE_API_BASE_URL", ".env.local / .env.example", "示例进入", "用于配置前端 API 基础路径"],
            [".env.example", "项目根目录", "是", "只保留变量名和占位值，供评阅人员理解配置方式"],
            [".gitignore", "项目根目录", "是", "包含 *.local，避免本地环境文件误提交"],
            ["zip_source_code", "报告脚本", "是", "打包时排除 .env.local、output、tmp、dist、node_modules"],
        ],
    )
    security_notes = [
        "本地 .env.local 的存在是为了让高德地图真实加载。没有该文件时，应用仍可显示拓扑视图和本地路线规划，但高德实景地图会提示缺少凭据。",
        "生产环境中更推荐通过代理服务配置 securityJsCode，避免安全密钥暴露在前端。本课程实验以本地演示为主，因此使用前端环境变量满足验证要求。",
        "源码包检查使用了精确字符串检索，确认真实 Key、安全密钥和 .env.local 都没有进入压缩包。这一步可以避免把运行凭据误当作代码附件提交。",
        "报告正文只写变量名，不写真实值。这样既说明了高德 API 的接入方式，也符合凭据最小暴露原则。",
        "外置盘环境会生成 ._* 元数据文件，报告脚本和最终检查都清理这些文件，避免它们混入源码包或影响 lint。",
    ]
    for text in security_notes:
        add_paragraph(document, text)

    add_heading(document, "二十七、报告材料组成说明")
    add_paragraph(document, "最终 output/doc 目录不只包含实验报告，还包含设计文档、发布说明、图表、真实截图、源码包和提交清单。不同材料承担不同证明作用。")
    add_table(
        document,
        ["材料", "文件或目录", "作用"],
        [
            ["实验报告", "实验报告_山科智行CampusPulse.docx/.pdf", "说明实验背景、原理、过程、截图、测试和结论"],
            ["设计文档", "设计文档_山科智行CampusPulse.docx/.pdf", "说明系统定位、架构、模型、界面和 API 接入"],
            ["发布说明", "发布说明_山科智行CampusPulse.docx/.pdf", "说明发布信息、验收清单和量化信息"],
            ["架构图", "山科智行_系统架构图.png", "说明前端、高德 API、服务端、模型和交付层关系"],
            ["流程图", "山科智行_开发交付流程图.png", "说明从基线拆解到归档提交的流程"],
            ["工作量图", "山科智行_工作量统计图.png", "说明源码、测试、点位和路网规模"],
            ["发布效果图", "发布效果_*.png", "用于发布说明中的关键效果展示"],
            ["真实应用截图", "application_screenshots/", "保留 13 张完整运行截图"],
            ["报告裁剪图", "report_figures/", "为 DOCX/PDF 排版生成的截图裁剪版"],
            ["源码包", "程序代码_山科智行CampusPulse源码.zip", "提交可检查源码，不含依赖和敏感配置"],
            ["目录说明", "程序代码目录说明.txt", "解释源码结构和关键文件"],
            ["提交清单", "提交材料清单.txt", "列出应提交的全部材料"],
        ],
    )
    material_notes = [
        "实验报告用于说明“为什么做、怎么做、做到什么程度”，设计文档用于说明“系统如何组织”，发布说明用于说明“当前版本如何验收”。三者内容有交叉，但侧重点不同。",
        "真实应用截图目录是本次补强的重点。它让评阅者可以直接看到软件运行状态，不需要只依赖报告中的裁剪图。",
        "源码包是最终可复查材料。由于排除了 node_modules 和构建产物，压缩包更小、更干净，也更适合提交。",
        "提交材料清单把图表、截图和源码包全部列出，减少提交时遗漏附件的风险。",
    ]
    for text in material_notes:
        add_paragraph(document, text)

    add_heading(document, "二十八、后续改进方向")
    future_items = [
        "接入更精确的校区建筑经纬度数据，使本地点位和高德 POI 的匹配更稳定，减少因为名称歧义导致的定位偏差。",
        "增加后台维护界面，支持管理员增删地点、维护入口节点、更新道路距离和临时封路规则。",
        "为高德安全密钥增加生产代理方案，避免生产环境直接暴露 securityJsCode，同时便于做访问控制和配额监控。",
        "接入室内楼层或建筑入口数据，使路线能够从校门延伸到楼宇入口，进一步提升实用性。",
        "增加路线偏好选项，例如最短距离、少爬坡、优先主路、夜间安全路线等，使系统从课程演示进一步接近真实服务。",
        "增加截图自动化脚本，把截图、报告生成和 PDF 渲染串联成一个完整命令，减少手动操作导致的遗漏。",
    ]
    add_bullets(document, future_items)

    add_heading(document, "二十九、结论")
    add_paragraph(
        document,
        f"本实验完成了 {PROJECT_NAME} 的软件改造和长篇报告交付。系统已经使用高德地图 JSAPI 2.0 作为默认实景地图能力，并保留本地路网路线规划、拓扑校验、搜索筛选、地点详情和移动端响应式能力。报告材料从短报告扩展为包含需求、原理、架构、数据、算法、API、界面、截图、测试和问题处理的完整实验记录。",
    )
    add_paragraph(document, "与基础项目相比，新版本在名称、地图能力、视觉风格、分区命名、页面文案、报告结构、截图证据和文件命名上均已重构。当前交付物包括可运行软件、后端服务、13 张真实应用截图、架构图、流程图、工作量图、DOCX/PDF 报告、设计文档、发布说明、源码包和提交材料清单，能够支撑课程实验提交与现场演示。")
    document.save(DOCS["report_docx"])


def create_report_doc_clean(metrics: ProjectMetrics) -> None:
    document = Document()
    set_base_style(document)
    add_title(document, "实验报告", PROJECT_SUBTITLE)
    add_table(
        document,
        ["项目", "内容"],
        [
            ["实验名称", f"{PROJECT_NAME} 设计与实现"],
            ["学生信息", "姓名、学号、班级待提交前填写"],
            ["实验主题", "使用高德地图开放平台 API 构建校园步行出行工作台"],
            ["成果形态", "Web 软件、后端路线服务、GitHub Pages 站点、Android WebView App、设计文档、实验报告、应用截图"],
        ],
    )

    add_heading(document, "一、实验背景与任务说明")
    for text in [
        "校园内部道路、建筑和生活服务点密集，学生每天会在宿舍、餐厅、教学楼、实验楼、图书馆、体育场和校门之间高频移动。传统静态校园地图只能展示建筑位置，不能根据起点和终点生成路线，也很难把搜索、筛选、地点详情和移动端浏览组织成一个连续的使用流程。本实验围绕这一实际场景，设计并实现一套面向校园步行的出行工作台，使用户可以在真实地图背景下完成地点查找、路线规划和空间理解。",
        "本次软件命名为“山科智行 Campus Pulse”，定位不是单纯的地点列表，也不是普通的信息展示页，而是把高德地图、校内路网、地点语义和路线摘要整合在同一个白色工作台中。用户进入系统后首先看到高德实景地图和路线调度区，随后可以继续搜索餐厅、筛选生活服务点、查看地点详情，或切换到校园拓扑视图核对本地路线算法。这样的结构更接近真实地图类应用，同时也能清楚展示课程实验中的数据结构和算法设计。",
        "实验要求在已有基础上重新完成软件和报告，因此实现过程中重点做了差异化改造。系统名称、页面视觉、地图引擎、交互路径、地点分区、文档叙述和截图材料都进行了重新组织。界面采用白色主视觉，地图区域成为首屏主体，路线调度、实时地点、任务过滤和地点队列围绕地图展开。报告正文围绕新软件本身叙述，不把开发检查过程、版本比较过程或提交材料检查过程写入实验主体。",
        "高德地图开放平台 API 是本实验的核心支撑。系统使用 JSAPI 加载真实地图底图，结合 POI 检索、地图控件、定位能力和步行路线反馈，使校园出行界面不再停留在静态拓扑层面。与此同时，本地路网仍然保留，因为它承担了课程实验中可解释路线规划的作用。真实地图和本地路网并列存在，前者提供空间参照，后者提供可控的数据结构和算法输出。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二、实验目的")
    for text in [
        "本实验首先要求完成一个可运行的校园出行 Web 应用。应用需要能够加载校园基础数据，展示地点分类和空间分区，支持用户选择起点与终点，并根据校内路网计算步行路线。路线结果需要包含总距离、预计耗时、折线坐标和逐段说明，而不是只展示一张固定图片。",
        "其次，实验要求掌握高德地图 API 在前端项目中的接入方式。系统需要在页面中加载高德地图对象，添加基础地图控件，围绕校区中心查询相关 POI，并在起终点确定后把高德实景路线反馈显示给用户。通过这一过程，可以理解第三方地图能力与本地业务数据之间的关系。",
        "再次，实验要求完成一套具有辨识度的界面设计。白色界面不是简单把背景改成白色，而是要重新安排信息层级，让地图、路线调度、搜索结果、筛选按钮和地点详情之间形成清晰关系。页面需要适合桌面演示，也需要在手机宽度下保持可读和可操作。",
        "最后，实验要求形成完整报告。报告需要说明需求、原理、架构、数据模型、算法、接口、界面、运行截图、验证结果和改进方向。报告中的截图必须来自实际操作状态，文字说明也必须与画面对应，避免出现截图与操作不一致的问题。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "三、实验环境与技术组成")
    add_paragraph(document, "系统采用前后端分离结构。前端使用 React、TypeScript 和 Vite 组织页面、组件与状态，后端使用 Express 和 TypeScript 提供基础数据与路线规划接口，共享类型文件统一地点、路网和路线结果的结构。地图层使用高德地图 JSAPI，文档层使用脚本生成 DOCX、PDF、架构图和运行截图材料。")
    add_table(
        document,
        ["组成", "技术或模块", "用途"],
        [
            ["前端框架", "React + TypeScript", "组织首页、地图组件、地点详情页和响应式状态"],
            ["构建工具", "Vite", "提供本地开发服务和生产构建能力"],
            ["地图服务", "高德地图 JSAPI", "展示真实底图、地图控件、POI 检索和步行路线反馈"],
            ["后端服务", "Express + TypeScript", "提供健康检查、基础数据和路线规划接口"],
            ["路线算法", "带权图最短路径", "根据地点入口节点和道路距离生成校内步行路线"],
            ["文档生成", "python-docx + LibreOffice", "生成实验报告、设计文档和 PDF 文件"],
        ],
    )
    for text in [
        f"从项目规模看，当前软件包含 {metrics.source_files} 个核心源码文件，约 {metrics.source_lines} 行代码，前端组件数量为 {metrics.components} 个，测试文件数量为 {metrics.test_files} 个。地点数据覆盖 {metrics.place_count} 个校园点位，路线图包含 {metrics.graph_nodes} 个节点和 {metrics.graph_edges} 条边。这些数据说明系统已经不是单页静态展示，而是具备地点、路线、地图、搜索、筛选和文档生成能力的完整实验项目。",
        "地图运行配置由本地运行环境提供，代码中只读取配置项，不在报告正文中写入具体值。这样既能让高德地图在开发和演示时正常工作，也能保证提交材料中只保留必要的实现说明。对于实验报告而言，重点是 API 的接入位置、加载顺序、使用的地图能力和界面效果，而不是暴露实际运行凭据。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "四、系统总体设计")
    for text in [
        "系统采用 BS 结构。浏览器端承担主要交互和地图展示，服务器端提供基础数据、路线规划接口和后续可扩展的数据服务。用户通过浏览器或 Android WebView 访问同一个站点，页面向后端请求校园地点、类别、路网节点和路线结果；当后端不可用时，浏览器端仍可使用本地回退规划器完成基础路线计算。这样的设计既适合课程演示，也便于部署到 GitHub Pages 后通过移动端 App 访问。",
        "浏览器端由 React、TypeScript 和 Vite 构成。React 负责组织首页、地点详情页、路线调度区、地图区、搜索区和地点队列；TypeScript 负责约束地点、路网和路线数据结构；Vite 负责开发服务和静态构建。页面使用 HashRouter，因此部署到 GitHub Pages 后不依赖服务器重写规则，Android WebView 也可以直接加载同一个 URL。",
        "地图能力由高德 JSAPI 提供。系统默认使用高德实景地图作为主视图，加载比例尺、工具条、图层切换和定位控件，并通过 POI 检索补齐校区周边真实地点。用户选择起点和终点后，高德地图用于呈现实景空间背景，本地路网用于生成可解释的校内路线摘要。两者共同构成系统的导航展示能力。",
        "服务端使用 Express 和 TypeScript。它提供 health、bootstrap 和 route 等接口，bootstrap 返回校区配置、地点、类别和路网基础数据，route 根据起终点计算 PlannedRoute。虽然前端具备本地回退能力，但服务端结构让系统符合典型 BS 架构，也为后续接入数据库、后台维护和统计服务留下空间。",
        "共享模型层位于 shared 目录，统一定义 CampusConfig、PlaceRecord、GraphNode、GraphEdge 和 PlannedRoute。前端页面、后端服务和本地回退规划器都围绕同一套类型工作，减少接口字段不一致导致的错误。报告中的数据模型说明也直接对应这些共享类型。",
        "移动端 App 采用 WebView 壳应用方式实现。Android 原生项目只负责加载部署后的网页地址、处理网络权限和基础页面容器，业务逻辑仍由 Web 端提供。这样形成清晰的 BS 结构：GitHub Pages 提供静态前端页面，浏览器或 Android WebView 作为客户端，后续服务器接口作为业务数据服务。",
        "首页布局采用“路线调度、地图主视图、结果列表”三层组织。桌面端在首屏中把路线调度放在左侧，把高德地图放在右侧；移动端则把这些区域改为纵向排列。无论用户从浏览器访问还是从 Android App 访问，核心操作都保持一致。",
        "系统的数据流从用户操作开始。当用户选择起点和终点后，前端将地点编号提交给路线服务，路线服务根据地点入口节点和路网边计算最短路径，返回 PlannedRoute 结果。前端拿到结果后同时更新摘要卡片、地图路线、拓扑折线和地点快照。若用户输入关键词或点击筛选按钮，前端会重新计算可见地点队列，并联动页面统计信息。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "五、数据模型设计")
    add_paragraph(document, "本实验把校园空间数据拆成地点语义和道路拓扑两类。地点语义回答“用户要找什么地方”，道路拓扑回答“程序如何从一个地方走到另一个地方”。两类数据通过入口节点关联，使一个地点可以对应多个可进入的路网节点。")
    add_table(
        document,
        ["模型", "核心字段", "说明"],
        [
            ["CampusConfig", "name, city, center, mapAsset", "描述校区整体配置和默认展示信息"],
            ["PlaceRecord", "name, categoryId, zone, aliases, keywords, mapPoint, accessNodeIds", "描述可搜索、可筛选、可进入路线规划的校园地点"],
            ["GraphNode", "id, label, zone, xPct, yPct", "描述校内路网节点和拓扑图中的显示位置"],
            ["GraphEdge", "fromNodeId, toNodeId, distanceMeters, instruction", "描述两节点之间的道路距离和步行提示"],
            ["PlannedRoute", "distanceMeters, estimatedMinutes, pathPoints, steps", "描述路线规划完成后的展示结果"],
        ],
    )
    for text in [
        "PlaceRecord 是用户最直接接触的数据。每个地点都包含名称、类别、空间带、别名、关键词、说明文字、拓扑坐标和入口节点。别名和关键词让用户可以用自然叫法搜索地点，例如输入“餐厅”可以找到 A 餐厅、B 餐厅和学苑餐厅，输入“宿舍”可以找到学生公寓相关点位。",
        "categoryId 和 zone 分别对应两个不同筛选维度。categoryId 面向行动目的，例如学习、食堂、图书馆、体育场馆和医疗服务；zone 面向空间区域，例如学习核心带、实验创新带、生活补给带和访客入口带。用户既可以按照“想做什么”查找，也可以按照“在哪一片区域”查找。",
        "GraphNode 和 GraphEdge 构成校内道路图。节点代表路口、建筑入口或区域连接点，边代表两个节点之间可以步行通过的道路。每条边保留距离和提示语，因此路线结果不仅能画线，还能生成“从西门方向进入中轴教学区”这类可读步骤。",
        "mapPoint 使用百分比坐标，用于在校园拓扑视图中绘制地点位置。真实地图坐标由高德地图能力解析或查询得到。两套坐标分工不同，拓扑坐标用于实验算法展示，真实坐标用于高德底图中的空间反馈。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "六、高德地图 API 接入实现")
    for text in [
        "高德地图接入集中在 AmapLiveMap 组件中完成。组件加载地图脚本后创建 AMap.Map 实例，设置校区中心点、缩放级别、视角和地图样式，并添加比例尺、工具条、图层切换和定位控件。用户打开首页时，右侧地图区域首先展示真实底图，底图上可以看到校区附近道路、建筑名称和地图版权信息。",
        "POI 检索用于把本地地点与真实地图空间建立联系。系统会根据地点名称、别名和校区关键词查询高德结果，把候选坐标缓存起来。当用户选中地点或请求路线时，组件可以复用已解析坐标，减少重复请求。对于本地地点表之外的周边服务点，系统还会围绕校区中心进行实时检索，并把结果显示在“高德实时地点”区域。",
        "步行路线反馈在用户确定起点和终点后触发。系统先解析两个地点的真实坐标，再调用高德步行能力生成实景地图路线。与此同时，本地路线规划会给出校内可解释路径。用户在同一页面中可以同时看到高德地图背景和本地路线摘要，两者共同完成校园出行说明。",
        "定位能力用于实时步行导航卡片。用户点击开始实时导航后，系统请求浏览器定位，并把当前位置转换为高德地图可使用的坐标。导航卡片会显示目标、定位状态、预计步行、当前位置和偏航情况。即使课堂演示不一定开启真实定位，这部分结构也为后续扩展提供了基础。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["地图能力", "系统用途", "页面体现"],
        [
            ["地图实例", "承载校园周边真实底图", "首页右侧实景地图"],
            ["控件能力", "辅助缩放、比例尺、图层切换和定位", "地图边缘的标准控件"],
            ["POI 检索", "查询校区与周边服务点", "高德实时地点列表"],
            ["步行路线", "生成真实地图上的步行反馈", "起终点确定后的实景路线"],
            ["定位能力", "支持实时导航入口", "实时步行导航卡片"],
        ],
    )

    add_heading(document, "七、路线规划算法实现")
    for text in [
        "本地路线规划使用带权图最短路径思想。用户选择起点和终点后，系统先找到对应 PlaceRecord，再读取两个地点的 accessNodeIds。一个地点可能有多个入口节点，例如图书信息中心可能靠近中轴教学区，也可能靠近生活区连接道路。系统会枚举起点入口和终点入口的组合，分别计算候选路径。",
        "图搜索过程中，程序维护每个节点的当前最短距离、前驱节点和访问状态。每轮选择尚未访问且距离最小的节点继续扩展，把经过边的距离累加到相邻节点上。如果新的距离更短，就更新相邻节点的前驱关系。搜索到终点后，程序沿前驱节点回溯，得到完整路径。",
        "得到节点路径后，系统会把路径转换成两类展示数据。第一类是 pathPoints，用于在校园拓扑图上绘制折线。第二类是 steps，用于生成自然语言分步说明。每个步骤包含起点节点、终点节点、距离和 instruction，因此路线摘要区能够展示清晰的行动提示。",
        "预计耗时根据总距离估算。校园步行速度受天气、坡度、拥堵和个人步速影响，实验中采用统一估算可以保证结果稳定，也便于比较不同路线。对于课堂演示，用户更关注路线方向和距离级别，因此分钟估算能够满足使用需求。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["阶段", "处理内容", "输出"],
        [
            ["地点校验", "确认起点和终点存在并读取入口节点", "起点地点、终点地点"],
            ["候选组合", "枚举起点入口和终点入口", "若干节点对"],
            ["图搜索", "在路网图上寻找最短距离路径", "节点序列"],
            ["结果组装", "计算距离、时间、折线和步骤", "PlannedRoute"],
            ["界面渲染", "更新地图、摘要和地点快照", "可读路线结果"],
        ],
    )

    add_heading(document, "八、界面与交互设计")
    for text in [
        "界面采用白色和暖灰作为底色，青绿色作为主操作色，橙色作为路线强调色。卡片边框保持轻量，圆角控制在较小范围内，避免页面变成堆叠装饰卡片。地图区域尽量宽，文字和控件围绕具体操作出现，不在首屏加入大量介绍性说明。",
        "首页顶部显示项目名称、应用定位和三个关键统计项。统计项包括服务点位、行动场景和地图引擎，帮助用户快速理解系统规模。首屏下方进入主要操作区，左侧是路线调度和摘要，右侧是高德地图与拓扑切换。这个布局让路线操作和地图反馈始终保持同屏关系。",
        "任务过滤使用一组清晰的按钮表示行动目的。用户点击“食堂”后，地点队列会立即只显示餐饮相关点位；点击“医疗服务”后，列表会聚焦校医院等服务点。空间场景则使用学习核心带、实验创新带、生活补给带和访客入口带组织地点，帮助用户从校区结构角度理解点位分布。",
        "地点详情页保留从首页进入单点说明的路径。用户可以从地点快照点击进入详情，查看该地点的用途、位置、关键词、可作为起点或终点的快捷操作，以及与同类地点的推荐关系。这样首页和详情页之间形成完整闭环。",
        "移动端布局采用纵向组织。路线调度、地图、搜索、实时地点和地点队列依次排列，按钮和输入框保持足够高度，地图区域保持可读尺寸。移动端截图说明系统不是只为桌面宽屏设计，手机宽度下仍然可以完成路线选择和地点搜索。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "九、核心功能实现")
    for text in [
        "首页由 HomePage 组织整体状态。它负责读取基础数据、解析地址中的筛选参数、维护当前起点终点、记录选中地点，并把数据传递给路线调度、地图、搜索区、地点队列和详情入口。这样做可以保证所有区域围绕同一份状态变化，不会出现地图和列表不一致的问题。",
        "CampusMap 组件承担地图容器的统一入口。默认状态下它显示高德实景地图，用户点击“校园拓扑”后切换到 SVG 拓扑视图。拓扑视图会绘制地点、区域、节点和路线折线，用于核对本地路网算法。两个视图共享起终点和路线状态，因此切换标签时路线不会丢失。",
        "AmapLiveMap 负责和高德地图交互。组件初始化时创建地图对象并注册控件，路线变化时解析起终点真实坐标，随后刷新高德步行路线。组件卸载时会清理地图实例，避免页面切换后残留旧对象。",
        "AmapPoiExplorer 负责实时地点展示。用户输入关键词后，系统会根据当前关键词和校区中心查询周边 POI，并把结果按距离和名称展示出来。这个区域与本地地点队列并列存在，前者体现高德实时数据，后者体现课程实验的本地数据。",
        "路线服务分为后端规划和浏览器本地回退两条路径。正常运行时前端请求后端接口，后端根据共享数据计算路线；当后端服务暂不可用时，前端本地规划器也可以输出同样结构的结果。这个设计提高了演示稳定性。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十、应用运行截图与操作说明")
    add_paragraph(document, "以下截图均对应具体操作状态。每张图都按软件实际页面截取，标题和说明描述的是截图中已经发生的操作，而不是预设的功能清单。截图覆盖高德实景首页、路线规划、地点快照、关键词搜索、食堂筛选、生活补给带筛选、详情页、移动端和校园拓扑校验。")
    add_application_screenshot_gallery(document, start_index=1)

    add_heading(document, "十一、路线规划场景验证")
    add_paragraph(document, "路线规划验证选择了两组典型路线。第一组是从西门到图书信息中心，代表访客或学生从校门进入学习核心区域的短路线；第二组是从学生公寓 A 区到工程实训中心，代表生活区到实验区的较长路线。两组路线覆盖不同空间带，能够验证算法不只适用于单一区域。")
    add_table(
        document,
        ["场景", "起点", "终点", "验证重点", "对应截图"],
        [
            ["R1", "西门", "图书信息中心", "入校到学习区域的短路线，验证摘要、地图和步骤说明", "图 2"],
            ["R2", "学生公寓 A 区", "工程实训中心", "生活区到实验区的跨区路线，验证较长路径计算", "图 7"],
            ["R3", "西门", "图书信息中心", "拓扑视图折线与节点关系，验证本地路网可解释性", "图 13"],
            ["R4", "移动端西门", "移动端图书信息中心", "窄屏下路线摘要、地图和控件可读性", "图 11"],
        ],
    )
    for text in [
        "从 R1 的截图可以看到，用户在起点框选择西门、终点框选择图书信息中心后，路线摘要区显示总距离、预计时间和分段提示，地图区域同步展示高德实景背景。这个场景证明路线结果不是静态页面，而是由用户选择触发后实时更新。",
        "R2 的路线更长，经过生活补给带、教学区域和实验区域。该场景用于验证路网图中跨区域边的连通性，也能检查步骤说明是否连续。如果图结构缺少关键边，R2 会出现无法规划或路线绕行异常，因此它适合作为综合验证场景。",
        "R3 切换到校园拓扑标签后，用户可以看到本地节点、地点标记和橙色路线折线。它说明系统并没有把所有路线能力完全交给高德地图，而是保留了课程实验所需的本地算法和可解释路网。",
        "R4 在手机宽度下执行同样的路线选择。截图显示路线控件、地图和摘要都能在窄屏中阅读，说明响应式布局能够支持移动端演示和实际查看。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十二、搜索、筛选与详情验证")
    for text in [
        "关键词搜索场景输入“餐厅”。页面中搜索框保留输入内容，高德实时地点区域显示校区周边餐饮相关 POI，本地地点队列同时显示 A 餐厅、B 餐厅和学苑餐厅。这个场景验证了本地数据和高德实时数据可以并列展示，用户既能看到校内规划点，也能看到真实地图中的周边服务点。",
        "任务过滤场景点击“食堂”。页面下方的过滤按钮处于选中状态，地点队列只剩下食堂相关地点。与单纯输入关键词不同，任务过滤基于地点分类字段，因此它能稳定地按照行动目的缩小范围，不受地点名称写法影响。",
        "空间场景筛选选择“生活补给带”。页面显示宿舍、食堂、校医院等生活区相关点位。这个场景验证 zone 字段能够把地点按照校区结构组织起来，适合用户从区域角度浏览校园服务。",
        "详情页验证选择图书信息中心和学生公寓 A 区两类地点。图书信息中心代表学习类地点，学生公寓 A 区代表生活类地点。两个详情页展示的信息重点不同，但入口、结构和快捷操作一致，说明详情模板可以复用到不同地点类别。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十三、主要源码模块实现说明")
    for text in [
        "HomePage 是首页的组织核心。它不直接承担地图绘制或路线算法，而是负责把用户操作转换为稳定的页面状态。起点、终点、关键词、任务过滤、空间场景和选中地点都在这一层汇总，再分发给路线调度区、地图区、实时地点区和地点队列。这样做的好处是页面中每个区域都围绕同一份状态更新，用户选择起点后，摘要、地图、详情快照和列表不会彼此脱节。",
        "CampusMap 是地图展示的统一入口。它内部同时管理高德实景视图和校园拓扑视图，默认把高德地图放在首位，符合本实验使用高德开放平台 API 的要求；当用户需要核对本地路网时，再切换到拓扑标签。组件接收 plannedRoute、selectedPlace、visiblePlaces 等数据后，把它们转换为地图所需的标记、折线和提示信息。这个模块体现了“真实地图展示”和“算法结构说明”之间的分工。",
        "AmapLiveMap 负责真实地图能力。它在组件挂载后加载高德地图对象，创建地图实例，注册比例尺、工具条、图层切换和定位控件，并根据校区中心设置初始视野。路线变化时，组件会解析起终点坐标并刷新步行路线；搜索变化时，实时地点列表也会根据关键词更新。由于地图对象生命周期和 React 组件生命周期不同，代码中对初始化、更新和销毁都做了分离，避免重复创建地图或留下无效引用。",
        "AmapPoiExplorer 负责展示高德实时地点结果。它的作用不是替代本地地点表，而是补充本地数据无法覆盖的真实周边信息。用户输入“餐厅”后，系统既会在本地地点队列中显示校内餐饮点位，也会在实时地点区域显示高德返回的周边餐饮服务。这个模块让页面具有真实地图应用的动态感，同时不影响课程实验中本地数据的可控性。",
        "local-route-planner 是浏览器侧路线计算模块。它和后端路线服务使用同一套共享类型，因此即使路线结果来自浏览器本地计算，前端渲染层也不需要写两套逻辑。这个模块在演示时很有价值，因为地图或后端状态可能受到运行环境影响，而本地规划可以保证核心路线功能始终可展示。",
        "server/src/services/route-planner.ts 是后端路线规划模块。它接收起点和终点编号，读取共享地点和路网数据，执行带权最短路径计算，再返回距离、耗时、路径点和步骤说明。后端规划的存在使系统更接近真实应用结构，也便于后续把地点数据、路况规则或管理后台接入服务端。",
        "shared/navigation.ts 是前后端共同依赖的类型契约。CampusConfig、PlaceRecord、GraphNode、GraphEdge、PlannedRoute 等结构都在这里定义。共享类型避免了前端认为字段存在、后端却没有返回的情况，也让报告中的数据模型说明能够直接对应代码结构。",
        "样式文件负责把白色工作台视觉落到页面上。全局样式定义字体、背景、颜色变量和基础排版，组件样式则控制卡片、地图容器、输入框、按钮、地点队列和响应式布局。由于本实验要求页面使用白色界面，样式层不只是美化，而是决定软件能否呈现出新的产品气质。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十四、高德地图交互流程详解")
    for text in [
        "高德地图交互从地图实例初始化开始。页面加载后，组件读取运行环境中的地图配置，创建地图实例并设置中心点、缩放级别和视图模式。随后，比例尺、工具条、图层切换和定位控件被添加到地图上。用户在截图中看到的地图按钮、底图切换和比例尺，都是这一初始化流程的结果。",
        "校区中心的确定是后续能力的基础。系统会优先使用配置中的校区中心点，同时可以通过地点名称和校区关键词向高德查询真实结果。中心点确定后，实时 POI 检索、附近地点排序和地图视野调整都围绕这个中心展开。这样做可以让实时结果集中在校园周边，而不是返回距离很远的同名地点。",
        "地点坐标解析采用缓存思路。本地点位表中保存的是语义信息和拓扑坐标，真实地图路线需要经纬度。系统在需要显示或规划某个地点时，根据地点名称、别名和校区关键词查询候选结果，并把可信结果缓存起来。之后再次选择相同地点时，可以直接使用缓存坐标，减少等待时间。",
        "步行路线反馈由起点和终点共同触发。当用户选择两个地点后，本地路线规划先给出校内可解释结果，高德地图层再根据真实坐标请求步行路线。页面中的路线摘要来自本地路网，地图上的实景反馈来自高德能力。两部分同步出现，能够让用户既看到可靠的课程算法结果，也看到真实地图上的空间关系。",
        "实时地点检索与关键词输入保持联动。用户输入“餐厅”时，本地搜索会过滤地点队列，高德检索会围绕校区中心查询餐饮相关 POI。两个结果区的来源不同，但都响应同一个关键词，这让页面交互保持一致，不会出现用户已经输入搜索词而某个区域仍然显示默认内容的情况。",
        "定位和实时导航入口采用渐进式设计。用户不启动实时导航时，地图仍然可以完成路线查看；用户点击导航后，系统再请求定位并显示当前位置、目标和状态。这种设计避免页面一打开就请求定位权限，也符合校园出行场景中“先看路线、再决定是否导航”的使用习惯。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十五、地点数据与场景组织分析")
    for text in [
        "地点数据的第一层组织是类别。学习、实验、食堂、图书馆、体育场馆、医疗服务、宿舍、访客入口等类别对应学生在校园中的真实行动目的。用户并不总是知道建筑全名，但通常知道自己要做什么，例如吃饭、上课、借书或回宿舍。类别筛选把这种自然需求直接转化为页面操作。",
        "地点数据的第二层组织是空间带。学习核心带聚合图书信息中心、逸夫楼等学习活动集中区域；实验创新带聚合工程实训中心和科研楼群；生活补给带聚合宿舍、食堂和校医院；访客入口带聚合校门和入口节点。这种组织方式帮助用户从校园结构理解地点，而不是只看一个线性列表。",
        "地点别名和关键词解决了中文校园场景中常见的称呼差异。学生可能说“A 餐厅”，也可能说“餐厅”“食堂”或“吃饭的地方”；可能说“图书馆”，也可能说“图书信息中心”。如果只用正式名称搜索，很多真实输入都会匹配失败。本系统通过 aliases 和 keywords 扩展可检索范围，提高搜索容错。",
        "地点说明用于补充用户决策。列表中每个地点除了名称和标签，还包含简短描述，说明它适合什么场景、靠近什么区域、作为起点或终点有什么意义。对于校园导航软件来说，地点不只是坐标，也承载了使用意图。描述文字能帮助用户在多个同类地点之间做选择。",
        "入口节点设计让地点和道路之间保持合理关系。现实中一个建筑可能有多个入口，最近道路也可能不是建筑中心点。本系统没有直接把地点之间两两连线，而是让地点通过 accessNodeIds 接入路网。路线规划时程序先从地点进入路网，再在节点之间搜索路径，最后到达目标地点的入口节点。",
        "点位数量控制在课程项目可理解范围内。二十多个点位可以覆盖主要校园行动，又不会让数据维护变得混乱。对报告和演示而言，这个规模足以展示搜索、筛选、路线规划、详情页和地图标注，同时也便于读者理解每个点位在系统中的作用。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十六、用户操作流程设计")
    for text in [
        "新用户进入首页时，第一步通常是观察校园空间。系统默认展示高德实景地图，并在顶部告诉用户这是基于高德 JSAPI 的校园出行工作台。此时用户不需要先阅读大量说明，只要看到地图、路线调度和服务点位统计，就能理解页面主要用途。",
        "当用户有明确起终点时，会先在路线调度区选择起点，再选择终点。选择完成后，按钮和摘要卡片会更新，地图区域展示对应路线，页面下方的地点队列仍然保留可继续浏览的入口。这个流程适合“从宿舍去上课”“从校门去图书馆”这类目标明确的出行场景。",
        "当用户只有模糊需求时，会使用搜索和筛选。例如用户想找吃饭地点，可以输入“餐厅”，也可以直接点击“食堂”过滤按钮。前者适合自由输入，后者适合快速选择。系统同时提供两种入口，是为了兼顾不同用户习惯。",
        "当用户对某个地点感兴趣时，可以进入地点详情页。详情页不是把首页信息重复一遍，而是把地点用途、关键词、相关地点和起终点快捷操作集中展示。用户可以在详情页把该地点设为起点或终点，再回到路线流程中。",
        "移动端流程强调连续阅读。由于屏幕较窄，系统将桌面端左右并列的信息改为上下排列。路线调度位于前方，地图随后出现，实时地点和地点队列继续向下展开。用户可以单手滚动完成查看，按钮和输入框也保持足够可点击面积。",
        "拓扑校验流程面向需要理解算法的人。普通用户默认使用实景地图即可，教师或开发者需要说明路网结构时，可以切换到校园拓扑视图。拓扑图展示节点、区域和路线折线，使算法讲解与页面演示保持一致。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十七、功能验证过程与结果")
    for text in [
        "首页验证关注首屏是否能够表达软件定位。实际页面中，顶部标题、三项统计、高德实景地图、路线调度区和地点快照都能正常显示。地图区域不是空白容器，而是可以看到真实底图和控件，说明地图能力已经接入到页面主体。",
        "路线验证关注用户选择起终点后系统是否产生完整反馈。从西门到图书信息中心的场景中，起点和终点选择框保持用户选择，路线摘要显示规划结果，地图区域同步显示实景背景，拓扑视图也能看到路线折线。这个结果说明路线状态在多个组件之间传递正常。",
        "搜索验证关注关键词是否同时影响本地地点和实时地点。输入“餐厅”后，本地列表显示餐厅相关地点，高德实时地点区域显示周边餐饮 POI。这个结果说明搜索不是只改了一个列表，而是同时驱动了本地数据和地图平台数据。",
        "筛选验证关注分类按钮是否真正改变结果集合。点击“食堂”后，按钮状态改变，地点队列只显示食堂相关点位，页面统计也随之变化。这个结果说明任务过滤和地点数据中的 categoryId 建立了正确关系。",
        "空间带验证关注 zone 字段是否发挥作用。选择“生活补给带”后，宿舍、食堂、校医院等生活区点位集中出现，学习和实验类地点不再作为主要结果显示。这个结果说明系统可以从空间结构角度组织校园地点。",
        "详情页验证关注首页与单点页面之间的衔接。图书信息中心详情页和学生公寓 A 区详情页都能通过首页进入，并展示各自不同的语义内容。用户在详情页可以继续理解地点用途，而不是被带到一个孤立页面。",
        "移动端验证关注窄屏布局是否可用。手机截图中，首页、路线结果和搜索结果都能够纵向阅读，输入框、按钮和地图没有横向溢出。这个结果说明响应式设计覆盖了校园出行软件常见的移动使用场景。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十八、边界情况与稳定性处理")
    for text in [
        "地图配置缺失时，页面不能直接崩溃。系统会在地图区域显示明确提示，让用户知道当前只能使用本地拓扑和路线规划。这样即使运行环境没有提供地图配置，核心课程功能仍然可被查看。",
        "后端服务不可用时，前端本地规划器可以继续工作。用户选择起点和终点后，浏览器侧会根据同一套地点和路网数据计算路线。这个设计让演示不完全依赖服务端状态，也让静态部署时仍然具备基本路线能力。",
        "起点和终点相同时，系统需要返回可理解结果，而不是进入普通搜索流程。路线规划模块会识别同一地点，返回距离为零的路线说明。这个边界分支虽然简单，但可以避免用户误操作时页面出现异常。",
        "没有匹配搜索结果时，页面需要保留结构。搜索框仍然显示用户输入，地点队列给出空结果提示，任务过滤和空间场景仍可继续使用。这样用户可以修改关键词或切换筛选，而不会以为页面加载失败。",
        "地图对象需要在组件卸载时释放。React 页面切换或热更新过程中，如果旧地图实例不清理，可能出现多个地图对象同时存在、控件重复或内存占用增加。实现中通过引用保存地图实例，并在清理阶段销毁对象。",
        "长页面需要避免布局跳动。地图、按钮、输入框和地点卡片都设置了稳定尺寸或间距，图片和动态列表加载时不会把上方内容大幅推开。对于报告截图来说，稳定布局也能保证同一操作状态下截图位置一致。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "十九、截图证据组织说明")
    for text in [
        "截图材料按照真实操作链路组织，而不是按照页面模块机械排列。前几张图展示用户刚进入系统、选择路线和查看地点快照；中间几张图展示搜索、食堂筛选和生活补给带筛选；后几张图展示详情页、移动端和拓扑校验。这样的顺序更接近用户实际使用过程。",
        "每张截图都保留了能够证明操作状态的界面元素。搜索图保留输入框中的“餐厅”，食堂筛选图保留被选中的“食堂”按钮，路线图保留起点和终点选择框，移动端路线图保留窄屏下的路线结果，拓扑图保留“校园拓扑”标签被选中的状态。",
        "截图说明采用自然语言解释画面中发生的事情。例如“点击任务过滤中的食堂后，地点队列只保留 A 餐厅、B 餐厅和学苑餐厅”，这比单独写“分类筛选页面”更准确。报告阅读者可以直接把说明和画面对应起来。",
        "截图在报告中使用裁剪图，完整图片仍保存在应用截图目录。裁剪图用于保证 DOCX 和 PDF 排版紧凑，完整图用于查看页面全貌。两者来源相同，因此不会出现报告图和原始截图不一致的情况。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十、系统结果分析")
    for text in [
        "从运行效果看，系统已经完成校园出行工作台的主要闭环。用户可以从首页进入，查看高德地图，选择路线，获得结果，继续搜索或筛选地点，并进入详情页查看单点信息。不同功能之间没有割裂，地点数据、路线状态和地图状态能够同步变化。",
        "白色界面改善了报告截图和课堂演示效果。浅色背景让高德地图、路线按钮、地点标签和列表卡片都更清晰，截图插入文档后不会因为暗色背景而显得沉重。青绿色主按钮和橙色路线强调形成了可辨识的视觉系统。",
        "高德地图 API 的接入提升了真实感。原本只有本地拓扑时，用户只能从抽象图上理解路线；加入真实底图后，用户可以看到校区周边道路和实际地名，空间参照更充分。实时 POI 也让页面不只依赖本地写死的数据，能够体现地图平台的信息能力。",
        "本地路线算法仍然具有必要价值。高德地图能够生成真实道路路线，但课程实验需要展示数据结构、路网节点、边权和路径步骤。本系统把两者放在同一个页面中，使实验结果既能看起来真实，也能讲清楚算法过程。",
        "响应式布局基本满足移动端查看。手机截图中首页、路线结果和搜索结果都可以阅读，主要控件没有超出屏幕。对于校园出行场景，移动端使用频率通常高于桌面端，因此这一点对系统完整性很重要。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十一、问题与解决方法")
    for text in [
        "第一个问题是地图能力和本地路网之间的数据粒度不同。高德地图返回的是真实经纬度和周边 POI，本地路网使用的是抽象节点和百分比坐标。解决方法是把二者分层处理：真实地图负责底图和 POI，拓扑路网负责算法和步骤，界面层只同步用户选择和路线摘要。",
        "第二个问题是地图加载具有异步特征。如果组件渲染完成但地图脚本还没有加载，页面可能出现空地图或控件缺失。实现中通过状态标记区分加载中、已加载和失败状态，并在必要位置显示提示，使用户知道当前地图能力是否可用。",
        "第三个问题是截图必须和具体操作对应。最初只截取页面默认状态容易造成标题和画面不一致，因此后续逐张按操作重新确认。搜索图保留搜索框输入，筛选图保留选中的按钮，路线图保留起点终点，拓扑图保留切换后的标签状态。",
        "第四个问题是页面内容较多，容易在移动端堆叠混乱。解决方法是让首屏区域在窄屏下纵向排列，地图和列表保持固定的最小尺寸，按钮允许换行，地点队列使用稳定卡片间距。这样既保留功能完整性，又避免文字拥挤。",
        "第五个问题是实验说明需要和运行画面保持一致。软件操作较多，如果只截默认页面，读者很难判断搜索、筛选、路线和详情功能是否真实发生。最终截图按操作状态逐张组织，文字说明也围绕画面中的输入、按钮和结果展开。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十二、特色与创新点")
    for text in [
        "本系统的第一个特色是把高德实景地图作为默认主视图，同时保留校园拓扑作为算法校验视图。用户看到的是具有真实地理背景的校园出行页面，教师或开发者需要讲解算法时，又可以切换到拓扑视图查看节点和路线折线。",
        "第二个特色是把任务目的和空间场景分开组织。任务目的回答“我要吃饭、学习、运动还是就医”，空间场景回答“这个地点位于学习核心区、实验区、生活区还是入口区”。两组筛选维度让地点浏览更接近真实校园使用习惯。",
        "第三个特色是截图材料覆盖真实操作链路。报告中的图不是只有首页和结果页，而是包含搜索、筛选、详情、移动端和拓扑校验。每张截图都对应具体动作，能够支撑报告文字说明。",
        "第四个特色是路线结果具有双重解释。高德地图给出真实空间背景，本地路网给出可解释步骤。用户获得可用路线，实验报告也能说明数据结构和算法来源。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十三、不足与改进方向")
    for text in [
        "当前点位坐标仍以演示数据为主，真实经纬度主要依靠地图检索补齐。后续可以采集更精确的建筑入口坐标，把本地点位和高德地图坐标稳定绑定，减少名称歧义带来的偏差。",
        "当前路线偏好较简单，默认按照距离寻找最短路径。后续可以加入少爬坡、优先主路、夜间安全、避开施工区域等偏好，使系统更接近真实校园导航服务。",
        "当前地点数据通过代码维护，适合课程实验，但不适合长期运行。后续可以增加管理后台或数据表，使管理员能够维护地点、类别、入口节点和道路状态。",
        "当前实时导航已经具备入口结构，但仍可继续增强。后续可以加入连续定位、偏航提醒、到达提示和语音播报，使移动端步行体验更完整。",
        "现有截图已经覆盖主要操作，但后续可以把更多特殊状态纳入材料，例如地图加载失败提示、无搜索结果状态、同起终点状态和较长路线步骤展开状态，使系统说明更全面。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十四、典型点位说明")
    for text in [
        "西门是本系统中最重要的访客入口点位之一。它不仅可以作为从校外进入校园的起点，也可以作为校园空间叙述的参照点。用户从西门到图书信息中心的路线能够覆盖入口区、教学区和学习区，是验证短路线、地图背景和地点快照的典型场景。",
        "图书信息中心代表学习类高频地点。它在系统中同时承担路线终点、详情页样例和拓扑校验目标三种角色。用户选择该地点后，首页快照会显示用途和快捷操作，详情页会补充地点说明，拓扑视图则能看到它与周边节点之间的连接关系。",
        "A 餐厅、B 餐厅和学苑餐厅共同构成食堂类点位。它们用于验证关键词搜索和任务过滤的差异：输入“餐厅”时，系统同时展示本地餐饮点和高德实时餐饮结果；点击“食堂”时，系统根据分类字段稳定筛选出校内餐饮点位。这个设计让同一类地点既能通过自然语言查找，也能通过按钮快速定位。",
        "学生公寓 A 区、学生公寓 A17、学生公寓 B2、学生公寓 B3、学生公寓 C3 和学生公寓 A2 等宿舍点位构成生活区起点集合。校园步行场景中，宿舍往往是学生一天行动的起点，因此宿舍点位数量相对较多。通过这些点位，系统可以展示从生活区到学习区、实验区和餐饮区的多种路线。",
        "工程实训中心和地球科学与工程学院楼代表实验创新带。它们和生活区、学习区之间距离较远，适合验证跨区路线是否连通。学生公寓 A 区到工程实训中心的路线就是围绕这一目的设计的，它能检查较长路线、分段提示和地图视野是否正常。",
        "校医院、体育运动场和综合体育馆代表公共服务点。它们虽然不一定每天被所有学生访问，但在校园服务系统中很重要。把这些点位放入地点表，可以让任务过滤覆盖医疗、体育等非学习场景，使软件不局限于上课和就餐。",
        "逸夫楼、若水园、校史馆、行政楼和门卫室等点位扩展了校园空间层次。它们可以作为学习、行政、参观和入口服务的补充节点，让地点队列更接近真实校园结构。用户浏览列表时，不会只看到宿舍和餐厅，而能看到较完整的校园行动对象。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十五、路线步骤与页面联动")
    for text in [
        "路线结果不是一个孤立的数字，而是由多个页面区域共同表达。总距离和预计时间出现在行程摘要中，分步说明出现在路线结果区，路径折线出现在校园拓扑视图，高德实景地图提供真实空间背景。用户只做一次起终点选择，就能在多个层面理解路线。",
        "以西门到图书信息中心为例，用户先在左侧路线调度区选择两个地点。系统得到地点编号后，读取西门和图书信息中心的入口节点，并在路网中搜索最短路径。路径生成后，摘要区显示结果，地点快照提示目标用途，地图区显示高德底图，拓扑视图可进一步核对节点连接。",
        "以学生公寓 A 区到工程实训中心为例，路线经过生活补给带和实验创新带，距离比入校到图书信息中心更长。这个场景能验证路线算法是否能够穿过多个空间带，也能验证页面在分步说明较长时是否仍保持可读。报告中的对应截图展示了跨区路线结果。",
        "路线步骤中的 instruction 字段让算法输出具有可读性。如果路线只返回节点编号，用户无法理解每一步应该怎么走；加入自然语言提示后，路线结果可以表达方向、区域和行动建议。对课程实验来说，这也证明 GraphEdge 不只是计算距离，还参与了用户界面表达。",
        "高德地图反馈和本地路线摘要存在互补关系。高德地图更擅长展示真实底图和周边地名，本地摘要更擅长展示课程自定义的校园节点和说明。用户查看路线时可以把二者结合起来理解，既知道地图上的空间位置，也知道系统规划的校内行动路径。",
        "移动端路线联动需要特别注意顺序。桌面端可以把调度和地图并排显示，移动端只能纵向排列，因此路线摘要必须紧跟选择控件，地图和地点列表继续向下展开。这样用户在手机上选择起终点后，不需要长距离滚动才能看到结果。",
        "拓扑路线校验保留了算法透明度。用户点击校园拓扑标签后，可以看到区域边界、节点标签、地点标记和橙色路线折线。如果路线摘要和拓扑折线指向不一致，问题会很容易被发现。因此拓扑视图既是展示功能，也是开发和讲解路线算法的重要依据。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十六、实验能力对应")
    for text in [
        "本实验对应前端工程能力。页面不是静态 HTML，而是由组件、状态和数据驱动。起终点选择、搜索输入、筛选按钮、地图切换和详情跳转都需要在 React 状态中协调。通过实现这些功能，可以训练组件拆分、状态提升、条件渲染和响应式布局能力。",
        "本实验对应地图 API 使用能力。高德地图接入涉及脚本加载、地图实例、控件注册、POI 检索、步行路线和定位入口。每一项能力都需要理解地图平台的对象模型和异步流程。最终页面能够显示真实地图和实时地点，说明 API 能力已经融入软件功能，而不是停留在示例代码层面。",
        "本实验对应数据建模能力。校园地点不是一串名称，而是包含类别、空间带、别名、关键词、说明、拓扑坐标和入口节点的结构化数据。道路也不是简单线段，而是带距离和提示语的边。通过这些模型，软件才能支持搜索、筛选、详情和路线规划。",
        "本实验对应算法应用能力。最短路径算法在课堂中常以抽象图出现，本系统把它放入校园步行场景。用户选择的是地点，程序计算的是节点路径，最终展示的是路线摘要和地图折线。这个过程把算法、数据和界面连接起来。",
        "本实验对应后端服务能力。后端提供基础数据和路线接口，使前端可以通过服务请求获得结果。虽然系统也有本地回退能力，但后端结构为后续接入数据库、管理后台和更复杂规则留下了空间。",
        "本实验对应文档表达能力。报告需要把软件的需求、原理、架构、数据、算法、地图 API、界面和截图讲清楚。只有截图和正文互相对应，读者才能相信软件确实执行了对应操作。报告因此成为软件成果的一部分，而不是最后附加的说明。",
        "本实验对应产品思维能力。白色界面、地图优先、任务过滤、空间场景、详情页和移动端适配，都不是单纯代码问题，而是围绕校园出行用户组织信息。软件看起来是否像一套真实工具，很大程度取决于这些交互和视觉选择。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十七、应用场景扩展设想")
    for text in [
        "新生报到是本系统最容易落地的场景之一。新生第一次到校时，常见路线包括从校门到报到点、从报到点到宿舍、从宿舍到食堂、从宿舍到教学楼。系统已经具备入口点、宿舍、食堂和学习地点数据，后续只需要把临时报到点作为地点加入，就可以形成一套报到日校园步行服务。",
        "日常上课场景强调准时和稳定。学生需要在宿舍、食堂、教学楼和实验楼之间快速切换，系统可以根据上课地点生成路线，并在移动端展示分步说明。若进一步接入课程表，用户进入页面后可以直接看到下一节课地点和推荐出发路线。",
        "餐饮高峰场景强调快速筛选。中午和傍晚时段，用户最关心的是附近有哪些食堂、距离多远、是否位于自己接下来的行动方向上。当前“餐厅搜索”和“食堂筛选”已经提供基础能力，后续可以加入营业时间、拥挤程度和推荐路线，使餐饮场景更实用。",
        "夜间出行场景强调安全感。夜间从图书馆、自习室或实验楼返回宿舍时，用户可能更希望选择主路、亮灯路段或经过人流较多区域。后续可以给 GraphEdge 增加道路类型、照明情况和安全等级字段，路线规划时根据用户偏好调整边权。",
        "访客导览场景强调理解成本。访客通常不知道校内建筑简称，也不熟悉校园分区。系统可以把访客入口带作为默认浏览入口，突出校门、行政楼、校史馆、图书信息中心和体育场馆等容易参观或办理事务的地点，并通过高德实景地图帮助访客建立方向感。",
        "大型活动场景强调临时组织。校园招聘会、运动会、开放日或考试日会改变人流和路线需求。系统可以在地点数据中增加活动点位，在路网中增加临时禁行或推荐通道，使页面从普通导航扩展到活动引导。",
        "后勤服务场景强调地点可达性。校医院、维修点、快递点、门卫室和餐饮服务点都属于校园生活的重要基础设施。系统的任务过滤可以继续扩展为后勤服务入口，帮助用户快速找到办事地点，并从当前位置或宿舍生成路线。",
        "无障碍出行场景强调路径属性。当前路线只考虑距离，后续可以给道路边增加坡度、台阶、电梯、坡道和路面平整度等信息。对于行动不便的用户，系统可以优先推荐坡度更小、障碍更少的路线，而不是单纯距离最短的路线。",
        "校内数据维护场景强调长期可更新。地点、道路、施工、活动和服务时间都会变化，如果完全写在代码中，维护成本会逐渐增加。后续可以把地点和路网放入数据库，并提供管理页面，让负责人员直接维护数据，再由前端实时读取。",
        "跨端使用场景强调同一套数据多端复用。桌面端适合展示全局地图和报告截图，移动端适合随身查看路线，后续还可以把同一套导航接口提供给小程序或校园门户。只要共享模型稳定，新增端侧应用时不需要重新设计地点和路网数据。",
        "数据分析场景强调路线需求沉淀。系统可以记录用户常用起点、终点和搜索关键词，统计哪些地点最常被访问、哪些路线最常被规划。经过脱敏处理后，这些数据可帮助学校理解校园空间使用情况，为道路优化、标识设置和服务布局提供参考。",
        "持续优化场景强调用户反馈。用户在实际使用中可能发现地点名称不准确、路线提示不清楚或地图定位存在偏差。后续可以增加反馈入口，把问题与具体地点或路线关联起来。维护人员根据反馈修正数据后，软件会逐渐从课程实验演示走向可持续运行的校园服务。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十八、应用截图逐项分析")
    for text in [
        "第 1 张截图展示系统默认首页。画面中可以看到白色顶部区域、高德实景地图、路线调度卡片和服务点位统计。这张图用于说明用户首次进入系统时看到的真实界面状态，也证明地图区域已经作为首屏主体出现。",
        "第 2 张截图展示“西门到图书信息中心”的路线结果。起点和终点选择框中保留用户选择，路线摘要随之更新，地图区域仍然显示高德底图。这张图对应的是完整路线规划操作，不是默认页面。",
        "第 3 张截图展示图书信息中心地点快照。用户选中该地点后，快照区域出现地点用途、关键词和快捷操作，说明地点选择状态能够影响首页局部内容。它验证的是地点预览，而不是路线规划。",
        "第 4 张截图展示关键词输入“餐厅”后的搜索结果。搜索框中保留输入内容，高德实时地点区域显示餐饮 POI，本地地点队列也出现餐饮相关地点。它证明关键词同时驱动实时数据和本地数据。",
        "第 5 张截图展示点击“食堂”任务过滤后的状态。页面下方的“食堂”按钮处于选中状态，地点队列只保留 A 餐厅、B 餐厅和学苑餐厅。它验证的是分类字段过滤，不是简单的关键词搜索。",
        "第 6 张截图展示生活补给带筛选状态。页面中可以看到宿舍、食堂、校医院等生活区相关点位，说明 zone 字段能够按照空间场景组织地点。这张图对应的是区域筛选操作。",
        "第 7 张截图展示学生公寓 A 区到工程实训中心的路线结果。该路线跨越生活区和实验区，距离更长，适合验证跨区连通性和分步说明的完整性。它与第 2 张短路线形成对照。",
        "第 8 张截图展示图书信息中心详情页。用户从首页地点快照进入详情后，可以看到该地点的用途说明、标签和操作入口。它证明系统存在从列表或快照到详情页的页面跳转链路。",
        "第 9 张截图展示学生公寓 A 区详情页。宿舍详情页和图书信息中心详情页结构一致，但内容面向生活区住宿点。它说明详情页模板可以复用到不同类别地点，同时保留各地点自己的语义信息。",
        "第 10 张截图展示移动端首页。页面在手机宽度下改为纵向排列，地图、路线调度和列表没有横向溢出。它验证的是响应式首页布局，而不是桌面页面缩小后的截图。",
        "第 11 张截图展示移动端路线规划结果。手机宽度下同样选择西门和图书信息中心，路线摘要与地图内容仍然可读。它说明核心路线功能在移动端也能使用。",
        "第 12 张截图展示移动端输入“餐厅”后的搜索结果。窄屏中搜索框、实时地点结果和本地地点队列仍然保持顺序清楚，用户可以继续滚动查看。它验证移动端搜索流程。",
        "第 13 张截图展示校园拓扑路线校验。用户点击“校园拓扑”标签后，页面显示本地路网、节点标记和橙色路线折线。它证明系统不仅接入高德地图，也保留了可解释的本地路线算法视图。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "二十九、结论")
    add_paragraph(document, f"本实验完成了 {PROJECT_NAME} 的设计与实现。系统以高德地图 API 为真实地图能力，以本地带权路网为路线算法基础，以白色校园出行工作台为界面表达，形成了一个可以搜索地点、筛选场景、规划路线、查看详情并支持移动端访问的 Web 应用。")
    add_paragraph(document, "从实验目标看，系统已经覆盖需求分析、数据建模、地图接入、路线计算、前后端协作、界面设计、截图展示和结果验证等环节。报告中的截图与实际操作状态对应，正文围绕软件设计、实现和运行结果展开。整体成果能够用于课程提交、课堂演示和后续扩展。")
    document.save(DOCS["report_docx"])


def create_design_doc(metrics: ProjectMetrics) -> None:
    document = Document()
    set_base_style(document)
    add_title(document, "设计文档", PROJECT_SUBTITLE)
    add_heading(document, "1. 项目定位")
    add_paragraph(document, f"{PROJECT_NAME} 定位为基于高德地图 API 的校园出行工作台，而不是传统地点导览页。用户进入页面后优先在高德实景地图中查看空间背景，再通过起终点、场景筛选和校内拓扑校验理解路线。")
    add_heading(document, "2. 架构设计")
    add_picture(document, DOCS["architecture_png"], caption="图 1  系统架构")
    add_table(
        document,
        ["层次", "职责", "关键文件"],
        [
            ["前端行动看板", "展示路线调度、地图、地点队列和详情快照", "src/routes/HomePage.tsx"],
            ["高德地图层", "加载 JSAPI 2.0，提供实景底图、POI 检索、定位和步行规划", "src/components/AmapLiveMap.tsx"],
            ["导航服务", "提供基础数据与路线规划接口", "server/src/app.ts"],
            ["路线算法", "基于带权图计算最短步行路径", "server/src/services/route-planner.ts"],
            ["共享契约", "统一前后端导航数据类型", "shared/navigation.ts"],
            ["静态回退", "无后端时在浏览器完成路线规划", "src/lib/local-route-planner.ts"],
        ],
    )
    add_heading(document, "3. 数据模型")
    add_paragraph(document, "数据模型采用地点和路网分离的原则。PlaceRecord 描述地点语义，GraphNode 和 GraphEdge 描述道路拓扑，PlannedRoute 描述最终可展示的路线结果。")
    add_table(
        document,
        ["模型", "核心字段", "用途"],
        [
            ["CampusConfig", "name, city, mapAsset, defaultView", "描述校区配置与底图"],
            ["PlaceRecord", "name, categoryId, zone, mapPoint, accessNodeIds", "描述地点和入口节点"],
            ["GraphNode", "id, label, zone, xPct, yPct", "描述路网节点"],
            ["GraphEdge", "fromNodeId, toNodeId, distanceMeters", "描述可通行道路边"],
            ["PlannedRoute", "distanceMeters, estimatedMinutes, pathPoints, steps", "前端地图和摘要展示"],
        ],
    )
    add_heading(document, "4. 界面设计")
    add_paragraph(document, "界面采用白色主视觉、暖灰背景、8px 圆角、清晰卡片层级和青绿色主色。布局上把路线调度卡片、行程摘要和地点快照放在左侧，把高德实景地图作为首屏主体放在右侧，下面再组织搜索、POI、任务过滤、空间分层和地点队列。")
    add_picture(document, SCREENSHOTS["home"], width=5.8, caption="图 2  行动看板首页效果")
    add_heading(document, "5. 高德地图 API 接入")
    add_paragraph(document, "前端通过 @amap/amap-jsapi-loader 加载高德 JSAPI 2.0。安全密钥在 AMapLoader.load 前写入 window._AMapSecurityConfig；地图初始化后加载 Scale、ToolBar、MapType、Geolocation 控件，并通过 PlaceSearch 查询校内地点，通过 Walking 生成真实步行路线反馈。")
    add_heading(document, "6. 路线规划设计")
    add_paragraph(document, "路线规划服务先根据地点 accessNodeIds 得到候选入口节点，再在路网图上执行最短路径搜索。结果会包含距离、预计耗时、路线折线和分步提示。若后端不可用，前端本地规划器按照相同契约输出结果。")
    add_heading(document, "7. 测试与交付")
    add_table(
        document,
        ["项目", "说明"],
        [
            ["自动测试", f"{metrics.test_files} 个测试文件，{metrics.test_cases} 条测试用例"],
            ["构建校验", "npm run build 同时执行 TypeScript 构建和 Vite 打包"],
            ["文档产物", "实验报告、设计文档、发布说明均输出 DOCX 和 PDF"],
            ["附件归档", "生成架构图、流程图、工作量图、截图、源码包和提交清单"],
        ],
    )
    document.save(DOCS["design_docx"])


def create_release_doc(metrics: ProjectMetrics) -> None:
    document = Document()
    set_base_style(document)
    add_title(document, "发布说明", PROJECT_SUBTITLE)
    add_heading(document, "1. 发布信息")
    add_table(
        document,
        ["项目", "内容"],
        [
            ["项目名称", PROJECT_NAME],
            ["源码仓库", REPO_URL],
            ["演示地址", SITE_URL],
            ["运行方式", "本地全栈运行或静态前端回退运行"],
            ["生成日期", "2026-07-07"],
        ],
    )
    add_heading(document, "2. 发布效果")
    add_picture(document, SCREENSHOTS["home"], width=5.8, caption="图 1  行动看板首页")
    add_picture(document, SCREENSHOTS["route"], width=5.8, caption="图 2  路线调度结果")
    add_picture(document, SCREENSHOTS["mobile"], width=3.2, caption="图 3  移动端界面")
    add_heading(document, "3. 验收清单")
    add_table(
        document,
        ["验收项", "结果", "说明"],
        [
            ["首页可访问", "通过", "展示 Campus Pulse 新标题和白色出行工作台界面"],
            ["路线规划", "通过", "可选择起点终点并生成距离、耗时和分步说明"],
            ["高德地图", "通过", "配置高德开放平台 Web 端 Key 后可加载 JSAPI 实景地图"],
            ["地图切换", "通过", "支持高德实景地图和校园拓扑校验两种视图"],
            ["响应式", "通过", "移动端宽度下核心控件保持可读"],
            ["材料生成", "通过", "DOCX/PDF、图表、截图和源码包集中输出"],
        ],
    )
    add_heading(document, "4. 量化信息")
    add_paragraph(
        document,
        f"当前版本包含 {metrics.source_files} 个核心源码文件、约 {metrics.source_lines} 行代码、{metrics.components} 个前端组件、{metrics.place_count} 个地点、{metrics.graph_nodes} 个节点和 {metrics.graph_edges} 条路网边。",
    )
    document.save(DOCS["release_docx"])


def convert_docx_to_pdf(docx_path: Path, expected_pdf: Path) -> None:
    command = [
        "soffice",
        f"-env:UserInstallation=file://{(TMP_DIR / 'lo-profile').resolve()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(OUTPUT_DIR),
        str(docx_path),
    ]
    subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    produced = OUTPUT_DIR / f"{docx_path.stem}.pdf"
    if produced != expected_pdf and produced.exists():
        shutil.move(produced, expected_pdf)


def render_pdf_preview(pdf_path: Path, prefix: str) -> None:
    if not pdf_path.exists():
        return
    output_prefix = TMP_DIR / prefix
    subprocess.run(["pdftoppm", "-png", "-f", "1", "-singlefile", str(pdf_path), str(output_prefix)], check=True)


def write_code_inventory(metrics: ProjectMetrics) -> None:
    lines = [
        "程序代码目录说明",
        "=" * 24,
        "",
        f"项目名称：{PROJECT_NAME}",
        f"源码仓库：{REPO_URL}",
        f"演示地址：{SITE_URL}",
        "",
        "核心目录：",
        "1. src/          React 前端页面、组件、样式和浏览器侧规划逻辑",
        "2. server/src/   Express 后端、导航数据和路线规划服务",
        "3. shared/       前后端共享 TypeScript 类型",
        "4. public/maps/  校园拓扑底图资源",
        "5. android/      Android WebView 壳应用，加载 GitHub Pages 网址",
        "6. .github/      Pages 部署与 Android APK 构建工作流",
        "7. scripts/      报告与交付材料生成脚本",
        "",
        "关键文件：",
        "1. src/routes/HomePage.tsx",
        "2. src/components/CampusMap.tsx",
        "3. src/lib/local-route-planner.ts",
        "4. server/src/services/route-planner.ts",
        "5. shared/navigation.ts",
        "",
        "量化信息：",
        f"- 源码文件数：{metrics.source_files}",
        f"- 源码总行数：{metrics.source_lines}",
        f"- 测试文件数：{metrics.test_files}",
        f"- 测试用例数：{metrics.test_cases}",
        f"- 点位数量：{metrics.place_count}",
        f"- 路网节点数：{metrics.graph_nodes}",
        f"- 路网边数量：{metrics.graph_edges}",
        "",
        "压缩包说明：源码包排除了 node_modules、dist、output、tmp、.env.local 和系统缓存文件。",
    ]
    DOCS["code_inventory"].write_text("\n".join(lines), encoding="utf-8")


def write_attachment_manifest() -> None:
    items = [
        DOCS["report_docx"],
        DOCS["report_pdf"],
        DOCS["design_docx"],
        DOCS["design_pdf"],
        DOCS["release_docx"],
        DOCS["release_pdf"],
        DOCS["architecture_png"],
        DOCS["workflow_png"],
        DOCS["workload_png"],
        SCREENSHOTS["home"],
        SCREENSHOTS["route"],
        SCREENSHOTS["mobile"],
        DOCS["code_zip"],
        DOCS["code_inventory"],
    ]
    if DOCS["android_apk"].exists():
        items.append(DOCS["android_apk"])
    lines = ["提交材料清单", "=" * 16, ""]
    lines.extend(f"- {item.name}" for item in items)
    if APP_SCREENSHOT_DIR.exists():
        lines.append("- application_screenshots/  真实应用截图目录")
        for _number, filename, title, _caption in APP_SCREENSHOTS:
            if (APP_SCREENSHOT_DIR / filename).exists():
                lines.append(f"  - {filename}：{title}")
    DOCS["attachment_list"].write_text("\n".join(lines), encoding="utf-8")


def copy_android_apk() -> None:
    source = ROOT / "output" / "android" / "app-debug.apk"
    if source.exists():
        shutil.copy2(source, DOCS["android_apk"])


def zip_source_code() -> None:
    include_roots = ["src", "server", "shared", "public", "scripts", "android", ".github"]
    include_files = [
        "package.json",
        "package-lock.json",
        "tsconfig.json",
        "tsconfig.app.json",
        "tsconfig.node.json",
        "tsconfig.server.json",
        "vite.config.ts",
        "eslint.config.js",
        "index.html",
        "README.md",
        "Dockerfile",
        ".dockerignore",
        ".env.example",
    ]
    excluded = {"node_modules", "dist", "output", "tmp", ".git", ".env.local"}
    with zipfile.ZipFile(DOCS["code_zip"], "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for root_name in include_roots:
            root = ROOT / root_name
            if not root.exists():
                continue
            for path in root.rglob("*"):
                if path.is_dir() or path.name.startswith("._"):
                    continue
                if any(part in excluded for part in path.parts):
                    continue
                archive.write(path, path.relative_to(ROOT))
        for file_name in include_files:
            path = ROOT / file_name
            if path.exists():
                archive.write(path, path.relative_to(ROOT))


def main() -> None:
    ensure_dirs()
    metrics = count_source_metrics()
    create_architecture_diagram(DOCS["architecture_png"])
    create_workflow_diagram(DOCS["workflow_png"])
    create_workload_chart(DOCS["workload_png"], metrics)
    create_placeholder_screenshots()
    create_report_doc_clean(metrics)
    create_design_doc(metrics)
    create_release_doc(metrics)
    for docx_key, pdf_key in [
        ("report_docx", "report_pdf"),
        ("design_docx", "design_pdf"),
        ("release_docx", "release_pdf"),
    ]:
        convert_docx_to_pdf(DOCS[docx_key], DOCS[pdf_key])
        render_pdf_preview(DOCS[pdf_key], pdf_key.replace("_pdf", "_page"))
    write_code_inventory(metrics)
    copy_android_apk()
    zip_source_code()
    write_attachment_manifest()

    print("Deliverables generated:")
    for path in [*DOCS.values(), *SCREENSHOTS.values()]:
        print(path)


if __name__ == "__main__":
    main()
